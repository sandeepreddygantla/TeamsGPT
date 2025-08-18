import os
import json
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass, asdict
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import re
from pathlib import Path
# Support both OpenAI and Azure OpenAI
from langchain_openai import ChatOpenAI, OpenAIEmbeddings, AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain.schema import HumanMessage, SystemMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging
import sqlite3
import faiss
from collections import defaultdict
import threading
import time
import hashlib
import httpx
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure tiktoken cache
tiktoken_cache_dir = os.path.abspath("tiktoken_cache")
os.environ["TIKTOKEN_CACHE_DIR"] = tiktoken_cache_dir

# Verify tiktoken cache exists (optional check)
if os.path.exists(tiktoken_cache_dir):
    expected_file = os.path.join(tiktoken_cache_dir, "9b5ad71b2ce5302211f9c61530b329a4922fc6a4")
    if not os.path.exists(expected_file):
        logging.warning("Tokenizer file missing in 'tiktoken_cache' directory - performance may be affected")
else:
    logging.warning("Tiktoken cache directory not found - will be created automatically on first use")
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from dotenv import load_dotenv

# Force reload of environment variables
load_dotenv(override=True)

# Ensure logs directory exists
os.makedirs('logs', exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/meeting_processor.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# API key will be checked when llm/embedding_model are used, not at module import
logger.info("Meeting processor module loading...")

project_id = "openai-meeting-processor"  # Simple project ID for personal use
tiktoken_cache_dir = os.path.abspath("tiktoken_cache")
os.environ["TIKTOKEN_CACHE_DIR"] = tiktoken_cache_dir

# --- Dummy Auth Function (for compatibility) ---
def get_access_token():
    """
    Get access token. Automatically detects environment and returns appropriate token.
    For OpenAI: Returns None (uses API key)
    For Azure: Returns Azure AD token
    """
    try:
        # Check if Azure environment variables are set
        azure_client_id = os.getenv("AZURE_CLIENT_ID")
        azure_client_secret = os.getenv("AZURE_CLIENT_SECRET")
        
        if azure_client_id and azure_client_secret:
            # Azure environment detected
            logger.info("Azure environment detected - getting Azure AD token")
            auth = "https://api.uhg.com/oauth2/token"
            scope = "https://api.uhg.com/.default"
            grant_type = "client_credentials"
            
            with httpx.Client() as client:
                body = {
                    "grant_type": grant_type,
                    "scope": scope,
                    "client_id": azure_client_id,
                    "client_secret": azure_client_secret
                }
                headers = {"Content-Type": "application/x-www-form-urlencoded"}
                response = client.post(auth, headers=headers, data=body, timeout=60)
                response.raise_for_status()
                return response.json()["access_token"]
        else:
            # OpenAI environment - no token needed
            logger.info("OpenAI environment detected - using API key authentication")
            return None
            
    except Exception as e:
        logger.error(f"Error getting access token: {e}")
        # Fallback to OpenAI mode
        return None

# Global variable to store current model selection
current_model_name = "gpt-5"  # Default model

# Available models configuration
AVAILABLE_MODELS = {
    "gpt-5": {
        "name": "GPT-5",
        "model": "gpt-4o",  # Map to closest OpenAI equivalent for testing
        "temperature": 0,
        "max_tokens": 16000,
        "description": "Most advanced model for complex reasoning"
    },
    "gpt-4.1": {
        "name": "GPT-4.1",
        "model": "gpt-4o-mini",  # Map to closest OpenAI equivalent for testing
        "temperature": 0,
        "max_tokens": 16000,
        "description": "Enhanced GPT-4 model for balanced performance"
    }
}

def get_current_model_config():
    """Get the current model configuration"""
    return AVAILABLE_MODELS.get(current_model_name, AVAILABLE_MODELS["gpt-5"])

def set_current_model(model_name: str):
    """Set the current model globally"""
    global current_model_name, llm
    if model_name in AVAILABLE_MODELS:
        current_model_name = model_name
        # Reinitialize LLM with new model
        try:
            access_token = get_access_token()
            llm = get_llm(access_token)
            logger.info(f"Successfully switched to model: {model_name}")
            return True
        except Exception as e:
            logger.error(f"Error switching to model {model_name}: {e}")
            return False
    else:
        logger.warning(f"Model {model_name} not available")
        return False

def get_current_model_name():
    """Get the current model name"""
    return current_model_name

# --- OpenAI LLM Client ---
def get_llm(access_token: str = None, model_name: str = None):
    """
    Get OpenAI LLM client. access_token parameter is kept for compatibility
    but not used since OpenAI uses API key authentication.
    model_name parameter allows overriding the current global model.
    """
    try:
        # Get fresh API key each time to avoid caching issues
        current_api_key = os.getenv("OPENAI_API_KEY")
        if not current_api_key:
            logger.warning("OPENAI_API_KEY environment variable not set")
            return None
        
        # Use provided model_name or current global model
        model_config = AVAILABLE_MODELS.get(model_name or current_model_name, AVAILABLE_MODELS["gpt-5"])
        
        return ChatOpenAI(
            model=model_config["model"],
            openai_api_key=current_api_key,
            temperature=model_config["temperature"],
            max_tokens=model_config["max_tokens"],
            request_timeout=120  # Increased timeout for complex queries
        )
    except Exception as e:
        logger.error(f"Error creating LLM client: {e}")
        return None

# --- OpenAI Embedding Model ---
def get_embedding_model(access_token: str = None):
    """
    Get OpenAI embedding model. access_token parameter is kept for compatibility
    but not used since OpenAI uses API key authentication.
    """
    try:
        # Get fresh API key each time to avoid caching issues
        current_api_key = os.getenv("OPENAI_API_KEY")
        if not current_api_key:
            logger.warning("OPENAI_API_KEY environment variable not set")
            return None
        
        return OpenAIEmbeddings(
            model="text-embedding-3-large",  # Using text-embedding-3-large
            openai_api_key=current_api_key,
            dimensions=3072  # text-embedding-3-large dimension
        )
    except Exception as e:
        logger.error(f"Error creating embedding model: {e}")
        return None

# Initialize global variables (keeping same structure)
# These will be None if API keys are not available, but won't crash the module load
try:
    access_token = get_access_token()
    embedding_model = get_embedding_model(access_token)
    llm = get_llm(access_token)
    logger.info("LLM and embedding model initialized successfully")
except Exception as e:
    logger.warning(f"Could not initialize LLM/embedding model at startup: {e}")
    access_token = None
    embedding_model = None
    llm = None

@dataclass
class DocumentChunk:
    """Structure to hold document chunk information with enhanced intelligence metadata"""
    chunk_id: str
    document_id: str
    filename: str
    chunk_index: int
    content: str
    start_char: int
    end_char: int
    embedding: Optional[np.ndarray] = None
    # Database metadata fields
    user_id: Optional[str] = None
    meeting_id: Optional[str] = None
    project_id: Optional[str] = None
    date: Optional[datetime] = None
    document_title: Optional[str] = None
    content_summary: Optional[str] = None
    main_topics: Optional[str] = None
    past_events: Optional[str] = None
    future_actions: Optional[str] = None
    participants: Optional[str] = None
    # Enhanced intelligence metadata
    enhanced_content: Optional[str] = None
    chunk_type: Optional[str] = None
    speakers: Optional[str] = None  # JSON string
    speaker_contributions: Optional[str] = None  # JSON string
    topics: Optional[str] = None  # JSON string
    decisions: Optional[str] = None  # JSON string
    actions: Optional[str] = None  # JSON string
    questions: Optional[str] = None  # JSON string
    context_before: Optional[str] = None
    context_after: Optional[str] = None
    key_phrases: Optional[str] = None  # JSON string
    importance_score: Optional[float] = None

@dataclass
class User:
    """Structure to hold user information"""
    user_id: str
    username: str
    email: str
    full_name: str
    password_hash: str
    created_at: datetime
    last_login: Optional[datetime] = None
    is_active: bool = True
    role: str = 'user'

@dataclass
class Project:
    """Structure to hold project information"""
    project_id: str
    user_id: str
    project_name: str
    description: str
    created_at: datetime
    is_active: bool = True

@dataclass
class Meeting:
    """Structure to hold meeting information"""
    meeting_id: str
    user_id: str
    project_id: str
    meeting_name: str
    meeting_date: datetime
    created_at: datetime

@dataclass
class MeetingDocument:
    """Structure to hold meeting document information"""
    document_id: str
    filename: str
    date: datetime
    title: str
    content: str
    content_summary: str  # Condensed summary for metadata
    main_topics: List[str]
    past_events: List[str]
    future_actions: List[str]
    participants: List[str]
    chunk_count: int = 0
    file_size: int = 0
    user_id: Optional[str] = None
    meeting_id: Optional[str] = None
    project_id: Optional[str] = None
    folder_path: Optional[str] = None
class EnhancedMeetingDocumentProcessor:
    """Enhanced Meeting Document Processor with Vector Database Support"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200, db_manager=None):
        """Initialize the enhanced processor
        
        Args:
            chunk_size: Size of text chunks for processing
            chunk_overlap: Overlap between chunks  
            db_manager: External database manager instance (if None, creates new one)
        """
        global llm, embedding_model, access_token
        self.llm = llm
        self.embedding_model = embedding_model
        self.access_token = access_token
        self.token_expiry = datetime.now() + timedelta(hours=1)
        
        # Text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""]
        )
        
        # Vector database - use provided manager or create new one
        if db_manager is not None:
            self.vector_db = db_manager
            logger.info("Using provided database manager")
        else:
            from src.database.manager import DatabaseManager
            self.vector_db = DatabaseManager()
            logger.info("Created new database manager")
        
        logger.info("Enhanced Meeting Document Processor initialized with OpenAI")
    
    def _detect_timeframe_from_query(self, query: str) -> Optional[str]:
        """Enhanced timeframe detection from natural language query"""
        query_lower = query.lower()
        
        # Comprehensive timeframe patterns with priority
        timeframe_patterns = [
            # Current periods (highest priority)
            (['current week', 'this week'], 'current_week'),
            (['current month', 'this month'], 'current_month'),
            (['current quarter', 'this quarter'], 'current_quarter'),
            (['current year', 'this year'], 'current_year'),
            
            # Last periods (high priority)
            (['last week', 'past week', 'previous week'], 'last_week'),
            (['last month', 'past month', 'previous month'], 'last_month'),
            (['last quarter', 'past quarter', 'previous quarter'], 'last_quarter'),
            (['last year', 'past year', 'previous year'], 'last_year'),
            
            # Specific day counts (medium priority)
            (['last 7 days', 'past 7 days', 'last seven days'], 'last_7_days'),
            (['last 14 days', 'past 14 days', 'last two weeks'], 'last_14_days'),
            (['last 30 days', 'past 30 days', 'last thirty days'], 'last_30_days'),
            (['last 60 days', 'past 60 days', 'last sixty days'], 'last_60_days'),
            (['last 90 days', 'past 90 days', 'last ninety days'], 'last_90_days'),
            
            # Extended periods (lower priority)
            (['last 3 months', 'past 3 months', 'last three months'], 'last_3_months'),
            (['last 6 months', 'past 6 months', 'last six months'], 'last_6_months'),
            (['last 12 months', 'past 12 months', 'last twelve months'], 'last_12_months'),
            
            # Recent periods (lowest priority)
            (['recent', 'recently', 'lately'], 'recent'),
        ]
        
        # Find the best match (earliest occurrence has highest priority)
        for patterns, timeframe in timeframe_patterns:
            for pattern in patterns:
                if pattern in query_lower:
                    return timeframe
        
        return None
    
    def _generate_date_based_summary(self, query: str, documents: List[Any], timeframe: str, include_context: bool = False) -> Union[str, Tuple[str, str]]:
        """Generate intelligent date-based summary with chronological organization"""
        
        # Sort documents by date
        sorted_docs = sorted(documents, key=lambda x: x.date)
        
        # Group documents by date for better organization
        from collections import defaultdict
        date_groups = defaultdict(list)
        for doc in sorted_docs:
            date_key = doc.date.strftime('%Y-%m-%d')
            date_groups[date_key].append(doc)
        
        # Build comprehensive context from all documents
        context_parts = []
        document_summaries = []
        
        for date_key, docs in sorted(date_groups.items()):
            date_formatted = datetime.strptime(date_key, '%Y-%m-%d').strftime('%B %d, %Y')
            context_parts.append(f"\n=== {date_formatted} ===")
            
            for doc in docs:
                # Add document summary
                doc_summary = f"Document: {doc.filename}\n"
                if doc.content_summary:
                    doc_summary += f"Summary: {doc.content_summary}\n"
                if doc.main_topics:
                    doc_summary += f"Main Topics: {', '.join(doc.main_topics)}\n"
                if doc.participants:
                    doc_summary += f"Participants: {', '.join(doc.participants)}\n"
                if doc.future_actions:
                    doc_summary += f"Action Items: {', '.join(doc.future_actions)}\n"
                
                context_parts.append(doc_summary)
                document_summaries.append(doc_summary)
        
        # Create comprehensive context
        full_context = '\n'.join(context_parts)
        
        # Generate summary prompt based on query type
        timeframe_display = timeframe.replace('_', ' ').title()
        summary_prompt = f"""
        The user asked: "{query}"
        
        Based on the meeting documents from {timeframe_display}, please answer their question naturally and comprehensively. Focus on what they specifically asked for rather than forcing a predetermined structure.

        Meeting Documents Context:
        {full_context}

        Provide a direct, helpful answer that addresses exactly what the user wants to know about {timeframe_display}.
        
        IMPORTANT: When referencing information, always cite the specific document filename rather than document numbers or chunk references.
        """
        
        try:
            # Use class LLM instance
            messages = [
                SystemMessage(content="You are a helpful AI assistant that answers questions about meeting documents naturally. Focus on what the user specifically asked for rather than forcing a predetermined structure. Always cite document filenames rather than chunk numbers when referencing information."),
                HumanMessage(content=summary_prompt)
            ]
            
            response = self.llm.invoke(messages)
            summary_response = response.content.strip()
            
            # Add timeframe context to the response
            final_response = f"**Summary for {timeframe_display}** ({len(documents)} documents)\n\n{summary_response}"
            
            if include_context:
                return final_response, full_context
            else:
                return final_response
                
        except Exception as e:
            logger.error(f"Error generating date-based summary: {e}")
            # Enhanced fallback using stored content summaries
            fallback_parts = [f"**Summary for {timeframe_display}** ({len(documents)} documents)\n"]
            
            for date_key, docs in sorted(date_groups.items()):
                date_formatted = datetime.strptime(date_key, '%Y-%m-%d').strftime('%B %d, %Y')
                fallback_parts.append(f"\n**{date_formatted}:**")
                
                for doc in docs:
                    fallback_parts.append(f"\n• **{doc.filename}**")
                    
                    # Include actual content summary if available
                    if doc.content_summary:
                        fallback_parts.append(f"  {doc.content_summary}")
                    
                    # Include main topics if available  
                    if doc.main_topics:
                        topics_str = ', '.join(doc.main_topics[:3])  # First 3 topics
                        fallback_parts.append(f"  *Topics: {topics_str}*")
                    
                    # Include participants if available
                    if doc.participants:
                        participants_str = ', '.join(doc.participants[:3])  # First 3 participants
                        fallback_parts.append(f"  *Participants: {participants_str}*")
            
            fallback_summary = '\n'.join(fallback_parts)
            
            if include_context:
                return fallback_summary, full_context
            else:
                return fallback_summary
    def refresh_clients(self):
        """Refresh OpenAI clients with new API key (if needed) and current model"""
        try:
            global access_token, llm, embedding_model

            if hasattr(self, 'token_expiry') and datetime.now() >= self.token_expiry - timedelta(minutes=5):
                logger.info("Refreshing access token...")
                access_token = get_access_token()
                self.access_token = access_token
                self.token_expiry = datetime.now() + timedelta(hours=1)

                # Use current model when refreshing
                llm = get_llm(access_token, current_model_name)
                embedding_model = get_embedding_model(access_token)
                self.llm = llm
                self.embedding_model = embedding_model
            else:
                # Force refresh even if not expired
                logger.info("Force refreshing access token...")
                access_token = get_access_token()
                self.access_token = access_token
                self.token_expiry = datetime.now() + timedelta(hours=1)

                # Use current model when refreshing
                llm = get_llm(access_token, current_model_name)
                embedding_model = get_embedding_model(access_token)
                self.llm = llm
                self.embedding_model = embedding_model

            logger.info(f"OpenAI clients refreshed successfully with model: {current_model_name}")
        except Exception as e:
            logger.error(f"Failed to refresh OpenAI clients: {e}")
            raise
    
    def extract_date_from_filename(self, filename: str, content: str = None) -> datetime:
        """
        Comprehensive date extraction with multiple methods and LLM validation:
        1. Try filename patterns first
        2. Try content extraction (top 10 lines)
        3. Try LLM extraction as fallback
        4. Validate final result with LLM
        """
        extracted_date = None
        extraction_method = "none"
        
        # 1. Try filename patterns first
        patterns = [
            r'(\d{8})_(\d{6})',  # YYYYMMDD_HHMMSS
            r'(\d{8})',          # YYYYMMDD
            r'(\d{4}-\d{2}-\d{2})', # YYYY-MM-DD
            r'(\d{2}/\d{2}/\d{4})'  # MM/DD/YYYY
        ]
        
        for pattern in patterns:
            match = re.search(pattern, filename)
            if match:
                try:
                    if '_' in match.group(0):
                        extracted_date = datetime.strptime(match.group(0), "%Y%m%d_%H%M%S")
                    elif '-' in match.group(0):
                        extracted_date = datetime.strptime(match.group(0), "%Y-%m-%d")
                    elif '/' in match.group(0):
                        extracted_date = datetime.strptime(match.group(0), "%m/%d/%Y")
                    else:
                        extracted_date = datetime.strptime(match.group(0), "%Y%m%d")
                    
                    extraction_method = "filename"
                    logger.info(f"Extracted date from filename: {extracted_date}")
                    break
                except ValueError:
                    continue
        
        # 2. Fallback: Extract date from file content if no filename date found
        if not extracted_date and content:
            try:
                extracted_date = self.extract_date_from_content(content)
                if extracted_date:
                    extraction_method = "content"
                    logger.info(f"Extracted date from content for {filename}: {extracted_date}")
            except Exception as e:
                logger.warning(f"Failed to extract date from content for {filename}: {e}")
        
        # 3. Validate the extracted date with LLM if we have both date and content
        if extracted_date and content:
            try:
                validated_date = self.validate_extracted_date(filename, extracted_date, content)
                if validated_date != extracted_date:
                    logger.info(f"LLM validation corrected date: {extracted_date} → {validated_date}")
                    extraction_method += "+validated"
                extracted_date = validated_date
            except Exception as e:
                logger.debug(f"Date validation failed: {e}")
        
        # 4. Final fallback: use current date if nothing worked
        if not extracted_date:
            logger.warning(f"Could not extract date from filename or content: {filename}, using current date")
            extracted_date = datetime.now()
            extraction_method = "fallback"
        else:
            logger.info(f"Final extracted date for {filename}: {extracted_date} (method: {extraction_method})")
        
        return extracted_date
    
    def extract_date_from_content(self, content: str) -> Optional[datetime]:
        """Extract date from meeting file content - checks top 10 lines comprehensively"""
        if not content:
            return None
            
        lines = content.strip().split('\n')
        
        # Check TOP 10 LINES for date (not just line 2)
        lines_to_check = lines[:10] if len(lines) >= 10 else lines
        
        logger.debug(f"Checking top {len(lines_to_check)} lines for date extraction")
        
        # Enhanced date patterns for meeting transcripts
        date_patterns = [
            # Full month names
            r'([A-Za-z]+\s+\d{1,2},\s+\d{4})',        # "July 14, 2025", "June 27, 2025"
            r'([A-Za-z]+\s+\d{1,2}\s+\d{4})',         # "July 14 2025" (no comma)
            
            # Numeric formats  
            r'(\d{1,2}/\d{1,2}/\d{4})',               # "7/14/2025", "07/14/2025"
            r'(\d{4}-\d{2}-\d{2})',                   # "2025-07-14"
            r'(\d{1,2}-\d{1,2}-\d{4})',               # "7-14-2025"
            r'(\d{1,2}\.\d{1,2}\.\d{4})',             # "7.14.2025" (European style)
            
            # Time included patterns (will extract just the date part)
            r'([A-Za-z]+\s+\d{1,2},\s+\d{4}),?\s+\d{1,2}:\d{2}', # "June 27, 2025, 2:30"
            r'(\d{1,2}/\d{1,2}/\d{4})\s+\d{1,2}:\d{2}',          # "6/27/2025 2:30"
        ]
        
        for line_num, line in enumerate(lines_to_check, 1):
            date_line = line.strip()
            
            if not date_line:
                continue
            
            logger.debug(f"Checking line {line_num}: {date_line[:100]}{'...' if len(date_line) > 100 else ''}")
            
            for pattern in date_patterns:
                match = re.search(pattern, date_line)
                if match:
                    date_str = match.group(1)
                    
                    # Clean up the date string (remove time, extra spaces, etc.)
                    date_str = re.sub(r',?\s+\d{1,2}:\d{2}.*$', '', date_str)  # Remove time
                    date_str = date_str.strip(' ,')
                    
                    logger.debug(f"Found potential date: '{date_str}' on line {line_num}")
                    
                    try:
                        # Enhanced parsing formats
                        parse_formats = [
                            "%B %d, %Y",    # "July 14, 2025"
                            "%B %d %Y",     # "July 14 2025"
                            "%b %d, %Y",    # "Jul 14, 2025" 
                            "%b %d %Y",     # "Jul 14 2025"
                            "%m/%d/%Y",     # "7/14/2025"
                            "%Y-%m-%d",     # "2025-07-14"
                            "%m-%d-%Y",     # "7-14-2025" 
                            "%m.%d.%Y",     # "7.14.2025"
                        ]
                        
                        for fmt in parse_formats:
                            try:
                                parsed_date = datetime.strptime(date_str, fmt)
                                logger.info(f"Successfully extracted date from content: {parsed_date} (line {line_num}, format: {fmt})")
                                return parsed_date
                            except ValueError:
                                continue
                                
                    except Exception as e:
                        logger.debug(f"Parse error for '{date_str}': {e}")
                        continue
        
        # If no standard patterns work, try using AI to extract the date
        logger.debug("Pattern matching failed, trying AI extraction")
        return self.ai_extract_date_from_content(content)
    
    def ai_extract_date_from_content(self, content: str) -> Optional[datetime]:
        """Use AI to extract meeting date from content when pattern matching fails"""
        if not self.llm:
            logger.debug("LLM not available for AI date extraction")
            return None
            
        try:
            # Get first 10 lines for AI analysis (same as pattern matching)
            first_lines = '\n'.join(content.strip().split('\n')[:10])
            
            extraction_prompt = f"""
Extract the meeting date from this document content. Look carefully at all the lines provided.

Content (first 10 lines):
{first_lines}

Instructions:
1. Look for any meeting date mentioned in the content
2. The date might be in various formats: "July 14, 2025", "7/14/2025", "2025-07-14", etc.
3. Ignore duration times like "31m 31s" - look for actual calendar dates
4. Return ONLY the date in YYYY-MM-DD format
5. If no clear meeting date is found, return "NONE"

Example responses:
- If you see "June 27, 2025, 2:30 PM" → return "2025-06-27"
- If you see "Meeting on 7/14/2025" → return "2025-07-14"  
- If no date found → return "NONE"

Return only the date in YYYY-MM-DD format or "NONE":
"""

            messages = [
                SystemMessage(content="You are a date extraction expert. Extract meeting dates accurately from document content and return them in YYYY-MM-DD format only. Be very careful to identify actual meeting dates, not durations or other time references."),
                HumanMessage(content=extraction_prompt)
            ]
            
            response = self.llm.invoke(messages)
            date_str = response.content.strip()
            
            logger.debug(f"AI date extraction response: '{date_str}'")
            
            # Validate the response format
            if re.match(r'^\d{4}-\d{2}-\d{2}$', date_str):
                parsed_date = datetime.strptime(date_str, "%Y-%m-%d")
                logger.info(f"AI successfully extracted date from content: {parsed_date}")
                return parsed_date
            else:
                logger.debug(f"AI date extraction returned invalid format or NONE: {date_str}")
                return None
                
        except Exception as e:
            logger.warning(f"AI date extraction failed: {e}")
            return None
    
    def validate_extracted_date(self, filename: str, extracted_date: datetime, content: str) -> datetime:
        """Use LLM to validate if the extracted date makes sense for this document"""
        if not self.llm or not content:
            return extracted_date  # Return as-is if no LLM or content
        
        try:
            # Get first 10 lines for validation context
            first_lines = '\n'.join(content.strip().split('\n')[:10])
            
            validation_prompt = f"""
I extracted a date of {extracted_date.strftime('%B %d, %Y')} from this document.

Document filename: {filename}
Document content (first 10 lines):
{first_lines}

Please validate:
1. Does this date make sense as the meeting date for this document?
2. Is there a different/better date mentioned in the content?

Respond with ONLY one of these formats:
- "VALID" - if the extracted date is correct
- "INVALID: YYYY-MM-DD" - if you found a better date (provide the better date)
- "INVALID: NONE" - if the date seems wrong but you can't find a better one

Examples:
- If extracted 2025-07-14 and content shows "July 14, 2025" → "VALID"
- If extracted 2025-07-14 but content shows "June 27, 2025" → "INVALID: 2025-06-27"  
- If extracted date seems wrong but no clear alternative → "INVALID: NONE"
"""

            messages = [
                SystemMessage(content="You are a meeting date validation expert. Validate whether extracted dates match the actual meeting date mentioned in document content."),
                HumanMessage(content=validation_prompt)
            ]
            
            response = self.llm.invoke(messages)
            result = response.content.strip()
            
            logger.debug(f"Date validation response: '{result}'")
            
            if result == "VALID":
                logger.debug("LLM confirmed extracted date is valid")
                return extracted_date
            elif result.startswith("INVALID: ") and result != "INVALID: NONE":
                # LLM found a better date
                better_date_str = result.replace("INVALID: ", "")
                if re.match(r'^\d{4}-\d{2}-\d{2}$', better_date_str):
                    better_date = datetime.strptime(better_date_str, "%Y-%m-%d")
                    logger.info(f"LLM suggested better date: {extracted_date} → {better_date}")
                    return better_date
            
            # Invalid but no better alternative, or other cases
            logger.debug("LLM validation failed or found no better alternative")
            return extracted_date
            
        except Exception as e:
            logger.debug(f"Date validation error: {e}")
            return extracted_date  # Return original on error
    
    def read_document_content(self, file_path: str) -> str:
        """Read document content from file"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
            elif file_ext == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                except ImportError:
                    logger.error("python-docx not installed. Cannot process .docx files.")
                    return ""
                    
            elif file_ext == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        content = ""
                        for page in pdf_reader.pages:
                            content += page.extract_text() + "\n"
                except ImportError:
                    logger.error("PyPDF2 not installed. Cannot process .pdf files.")
                    return ""
            else:
                logger.error(f"Unsupported file format: {file_ext}")
                return ""
            
            return content if content.strip() else ""
                
        except Exception as e:
            logger.error(f"Error reading document {file_path}: {e}")
            return ""
    
    def create_content_summary(self, content: str, max_length: int = 1500) -> str:
        """Create a condensed summary of the content"""
        try:
            summary_prompt = f"""
            Create a concise summary of this meeting document in 2-3 sentences (max {max_length} characters).
            Focus on the main purpose, key decisions, and outcomes.
            
            Content: {content[:2000]}...
            
            Summary:
            """
            
            messages = [
                SystemMessage(content="You are a meeting summarization expert. Create concise, informative summaries."),
                HumanMessage(content=summary_prompt)
            ]
            
            response = self.llm.invoke(messages)
            summary = response.content.strip()
            
            return summary[:max_length] if len(summary) > max_length else summary
            
        except Exception as e:
            logger.error(f"Error creating content summary: {e}")
            # Fallback to truncated content
            return content[:max_length] + "..." if len(content) > max_length else content
    
    def parse_document_content(self, content: str, filename: str, user_id: str, project_id: str = None, meeting_id: str = None) -> MeetingDocument:
        """Parse a meeting document and extract structured information"""
        doc_date = self.extract_date_from_filename(filename, content)
        document_id = f"{filename}_{doc_date.strftime('%Y%m%d_%H%M%S')}"
        
        parsing_prompt = f"""
        You are an expert document analyst. Analyze this meeting document and extract information in valid JSON format.
        
        Document content:
        {content[:8000]}{"..." if len(content) > 8000 else ""}
        
        Extract:
        1. "title": Clear, descriptive meeting title (2-8 words)
        2. "main_topics": Array of main topics discussed
        3. "past_events": Array of past events/completed items mentioned
        4. "future_actions": Array of upcoming actions/planned activities
        5. "participants": Array of participant names mentioned
        
        Return only valid JSON with no additional text.
        """
        
        try:
            messages = [
                SystemMessage(content="You are a document parsing assistant. Always return valid JSON format with no additional text."),
                HumanMessage(content=parsing_prompt)
            ]
            
            response = self.llm.invoke(messages)
            content_str = response.content.strip()
            
            # Clean JSON response
            if content_str.startswith('```json'):
                content_str = content_str[7:-3].strip()
            elif content_str.startswith('```'):
                content_str = content_str[3:-3].strip()
            
            parsed_data = json.loads(content_str)
            
            # Create content summary
            content_summary = self.create_content_summary(content)
            
            return MeetingDocument(
                document_id=document_id,
                filename=filename,
                date=doc_date,
                title=parsed_data.get('title', f'Meeting - {doc_date.strftime("%Y-%m-%d")}'),
                content=content,
                content_summary=content_summary,
                main_topics=parsed_data.get('main_topics', []),
                past_events=parsed_data.get('past_events', []),
                future_actions=parsed_data.get('future_actions', []),
                participants=parsed_data.get('participants', []),
                file_size=len(content),
                user_id=user_id,
                project_id=project_id,
                meeting_id=meeting_id
            )
            
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {e}")
            return self._create_fallback_document(content, filename, doc_date, document_id, user_id, project_id, meeting_id)
    
    def _create_fallback_document(self, content: str, filename: str, doc_date: datetime, document_id: str, user_id: str, project_id: str = None, meeting_id: str = None) -> MeetingDocument:
        """Create a fallback document when parsing fails"""
        content_summary = self.create_content_summary(content)
        
        return MeetingDocument(
            document_id=document_id,
            filename=filename,
            date=doc_date,
            title=f"Meeting - {doc_date.strftime('%Y-%m-%d')}",
            content=content,
            content_summary=content_summary,
            main_topics=[],
            past_events=[],
            future_actions=[],
            participants=[],
            file_size=len(content),
            user_id=user_id,
            project_id=project_id,
            meeting_id=meeting_id
        )
    
    def extract_meeting_intelligence(self, content: str, max_tokens: int = 115000) -> Dict:
        """Extract comprehensive meeting intelligence using LLM"""
        try:
            # Estimate content tokens and handle large meetings
            estimated_tokens = len(content.split()) * 1.3  # Rough estimation
            
            if estimated_tokens > max_tokens:
                # For large meetings, process in windows
                return self._extract_intelligence_windowed(content, max_tokens)
            else:
                # Single pass for smaller meetings
                return self._extract_intelligence_single_pass(content)
                
        except Exception as e:
            logger.error(f"Error extracting meeting intelligence: {e}")
            return self._create_fallback_intelligence()
    
    def _extract_intelligence_single_pass(self, content: str) -> Dict:
        """Extract intelligence in single LLM call"""
        intelligence_prompt = f"""
        COMPREHENSIVE MEETING INTELLIGENCE EXTRACTION
        
        Meeting Content:
        {content}
        
        Extract ALL information and structure as JSON:
        
        {{
            "meeting_metadata": {{
                "meeting_type": "standup/planning/demo/review/discussion",
                "main_purpose": "brief description of meeting purpose",
                "duration_estimate": "estimated duration if mentioned",
                "meeting_effectiveness": "productive/average/ineffective"
            }},
            "participants": [
                {{
                    "name": "Full Name",
                    "role": "role if mentioned",
                    "contribution_summary": "what they primarily discussed",
                    "speaking_frequency": "high/medium/low",
                    "key_statements": ["important quotes"],
                    "questions_asked": ["questions they posed"],
                    "decisions_influenced": ["decisions they affected"],
                    "expertise_demonstrated": ["areas of knowledge shown"]
                }}
            ],
            "topics_discussed": [
                {{
                    "topic": "topic name",
                    "discussed_by": ["speaker1", "speaker2"],
                    "key_points": ["point1", "point2"],
                    "outcome": "decision/action/discussion",
                    "outcome_details": "specific outcome",
                    "importance_score": 0.8
                }}
            ],
            "decisions_made": [
                {{
                    "decision": "what was decided",
                    "decided_by": "who made decision",
                    "context": "why this decision",
                    "impact": "what this affects",
                    "implementation_notes": "how to implement"
                }}
            ],
            "action_items": [
                {{
                    "task": "what needs to be done",
                    "assigned_to": "person responsible",
                    "due_date": "if mentioned",
                    "priority": "high/medium/low",
                    "dependencies": ["other tasks"],
                    "context": "why this task is needed"
                }}
            ],
            "questions_and_answers": [
                {{
                    "question": "what was asked",
                    "asked_by": "who asked",
                    "answered_by": "who responded",
                    "answer": "response given",
                    "resolved": true
                }}
            ],
            "meeting_flow": {{
                "opening_discussion": "how meeting started",
                "main_discussion_points": ["point1", "point2"],
                "key_transitions": ["transition descriptions"],
                "closing_notes": "how meeting ended"
            }},
            "contextual_references": {{
                "previous_meetings_mentioned": ["meeting references"],
                "external_projects_mentioned": ["project names"],
                "systems_discussed": ["system names"],
                "deadlines_mentioned": ["deadline references"]
            }},
            "searchable_metadata": {{
                "all_names_mentioned": ["name1", "name2"],
                "technical_terms": ["term1", "term2"], 
                "project_names": ["project1", "project2"],
                "key_phrases": ["phrase1", "phrase2"],
                "sentiment_tone": "positive/neutral/negative",
                "urgency_level": "high/medium/low"
            }}
        }}
        
        Be extremely thorough. Extract every piece of information that could be useful for future queries.
        Include speaker attributions for all content. Preserve context and nuance.
        """
        
        try:
            messages = [
                SystemMessage(content="You are an expert meeting intelligence analyst. Always return valid JSON with comprehensive information extraction."),
                HumanMessage(content=intelligence_prompt)
            ]
            
            response = self.llm.invoke(messages)
            content_str = response.content.strip()
            
            # Clean JSON response
            if content_str.startswith('```json'):
                content_str = content_str[7:-3].strip()
            elif content_str.startswith('```'):
                content_str = content_str[3:-3].strip()
            
            intelligence_data = json.loads(content_str)
            logger.info("Successfully extracted meeting intelligence")
            return intelligence_data
            
        except Exception as e:
            logger.error(f"Error in single-pass intelligence extraction: {e}")
            return self._create_fallback_intelligence()
    
    def _extract_intelligence_windowed(self, content: str, max_tokens: int) -> Dict:
        """Extract intelligence from large meetings using windowed approach"""
        # For now, use truncated content for large meetings
        # TODO: Implement full windowed processing if needed
        truncated_content = content[:max_tokens * 3]  # Rough token-to-char conversion
        logger.warning(f"Content truncated for intelligence extraction: {len(content)} -> {len(truncated_content)} characters")
        return self._extract_intelligence_single_pass(truncated_content)
    
    def _create_fallback_intelligence(self) -> Dict:
        """Create fallback intelligence structure when extraction fails"""
        return {
            "meeting_metadata": {
                "meeting_type": "discussion",
                "main_purpose": "General meeting discussion",
                "duration_estimate": "unknown",
                "meeting_effectiveness": "average"
            },
            "participants": [],
            "topics_discussed": [],
            "decisions_made": [],
            "action_items": [],
            "questions_and_answers": [],
            "meeting_flow": {
                "opening_discussion": "Meeting discussion",
                "main_discussion_points": [],
                "key_transitions": [],
                "closing_notes": "Meeting concluded"
            },
            "contextual_references": {
                "previous_meetings_mentioned": [],
                "external_projects_mentioned": [],
                "systems_discussed": [],
                "deadlines_mentioned": []
            },
            "searchable_metadata": {
                "all_names_mentioned": [],
                "technical_terms": [],
                "project_names": [],
                "key_phrases": [],
                "sentiment_tone": "neutral",
                "urgency_level": "medium"
            }
        }
    
    def chunk_document(self, document: MeetingDocument, intelligence_data: Dict = None) -> List[DocumentChunk]:
        """Split document into intelligent chunks with enhanced metadata"""
        # Extract intelligence if not provided
        if intelligence_data is None:
            intelligence_data = self.extract_meeting_intelligence(document.content)
        
        # Create intelligent chunks with context preservation
        intelligent_chunks = self._create_intelligent_chunks(document, intelligence_data)
        
        document.chunk_count = len(intelligent_chunks)
        return intelligent_chunks
    
    def _create_intelligent_chunks(self, document: MeetingDocument, intelligence_data: Dict) -> List[DocumentChunk]:
        """Create intelligent chunks with rich metadata"""
        # Split content into base chunks
        base_chunks = self.text_splitter.split_text(document.content)
        
        document_chunks = []
        current_pos = 0
        
        for i, chunk_content in enumerate(base_chunks):
            chunk_id = f"{document.document_id}_chunk_{i}"
            
            # Find start position in original content
            start_char = document.content.find(chunk_content, current_pos)
            if start_char == -1:
                start_char = current_pos
            
            end_char = start_char + len(chunk_content)
            current_pos = end_char
            
            # Generate embedding for chunk
            try:
                if self.embedding_model is None:
                    logger.warning("Embedding model not available, using zero vector")
                    embedding_array = np.zeros(3072)
                else:
                    embedding = self.embedding_model.embed_query(chunk_content)
                    embedding_array = np.array(embedding)
            except Exception as e:
                logger.error(f"Error generating embedding for chunk {chunk_id}: {e}")
                embedding_array = np.zeros(3072)
            
            # Extract chunk-specific intelligence
            chunk_intelligence = self._extract_chunk_intelligence(chunk_content, intelligence_data, i, len(base_chunks))
            
            # Create enhanced chunk with metadata
            chunk = DocumentChunk(
                chunk_id=chunk_id,
                document_id=document.document_id,
                filename=document.filename,
                chunk_index=i,
                content=chunk_content,
                start_char=start_char,
                end_char=end_char,
                embedding=embedding_array,
                # Copy essential metadata from document - FIX FOR USER_ID BUG
                user_id=document.user_id,
                meeting_id=document.meeting_id,
                project_id=document.project_id,
                date=document.date,
                document_title=document.title
            )
            
            # Add intelligence metadata
            chunk.enhanced_content = self._create_enhanced_content(chunk_content, chunk_intelligence)
            chunk.chunk_type = chunk_intelligence.get('chunk_type', 'discussion')
            chunk.speakers = json.dumps(chunk_intelligence.get('speakers', []))
            chunk.speaker_contributions = json.dumps(chunk_intelligence.get('speaker_contributions', {}))
            chunk.topics = json.dumps(chunk_intelligence.get('topics', []))
            chunk.decisions = json.dumps(chunk_intelligence.get('decisions', []))
            chunk.actions = json.dumps(chunk_intelligence.get('actions', []))
            chunk.questions = json.dumps(chunk_intelligence.get('questions', []))
            chunk.context_before = chunk_intelligence.get('context_before', '')
            chunk.context_after = chunk_intelligence.get('context_after', '')
            chunk.key_phrases = json.dumps(chunk_intelligence.get('key_phrases', []))
            chunk.importance_score = chunk_intelligence.get('importance_score', 0.5)
            
            document_chunks.append(chunk)
        
        return document_chunks
    
    def _extract_chunk_intelligence(self, chunk_content: str, meeting_intelligence: Dict, chunk_index: int, total_chunks: int) -> Dict:
        """Extract intelligence metadata for a specific chunk"""
        chunk_intel = {
            'chunk_type': 'discussion',
            'speakers': [],
            'speaker_contributions': {},
            'topics': [],
            'decisions': [],
            'actions': [],
            'questions': [],
            'context_before': '',
            'context_after': '',
            'key_phrases': [],
            'importance_score': 0.5
        }
        
        try:
            # Extract speakers mentioned in this chunk
            all_participants = meeting_intelligence.get('participants', [])
            chunk_speakers = []
            chunk_speaker_contributions = {}
            
            logger.info(f"Processing chunk {chunk_index}: Found {len(all_participants)} participants in intelligence data")
            if all_participants:
                logger.info(f"Participants: {[p.get('name', 'Unknown') for p in all_participants]}")
            
            for participant in all_participants:
                participant_name = participant.get('name', '')
                if participant_name:
                    # Enhanced name matching to handle different formats
                    name_parts = participant_name.split()
                    chunk_lower = chunk_content.lower()
                    
                    # Check multiple matching strategies
                    matched = False
                    
                    # Strategy 1: Exact name match
                    if participant_name.lower() in chunk_lower:
                        matched = True
                    
                    # Strategy 2: Last name, First name format (common in meeting transcripts)
                    elif len(name_parts) >= 2:
                        # Try "Last, First" format
                        last_first = f"{name_parts[-1]}, {' '.join(name_parts[:-1])}"
                        if last_first.lower() in chunk_lower:
                            matched = True
                        
                        # Try individual name parts
                        elif any(part.lower() in chunk_lower for part in name_parts if len(part) > 2):
                            matched = True
                    
                    # Strategy 3: Check if any significant name part appears
                    elif any(part.lower() in chunk_lower for part in name_parts if len(part) > 2):
                        matched = True
                    
                    if matched:
                        chunk_speakers.append(participant_name)
                        chunk_speaker_contributions[participant_name] = {
                            'contribution_summary': participant.get('contribution_summary', ''),
                            'speaking_frequency': participant.get('speaking_frequency', 'low')
                        }
                        logger.info(f"Found speaker {participant_name} in chunk {chunk_index}")
            
            if not chunk_speakers and all_participants:
                logger.info(f"No speakers matched in chunk {chunk_index}. First 200 chars: {chunk_content[:200]}")
            
            # Extract topics relevant to this chunk
            chunk_topics = []
            for topic in meeting_intelligence.get('topics_discussed', []):
                topic_name = topic.get('topic', '')
                if any(keyword.lower() in chunk_content.lower() for keyword in topic.get('key_points', [])):
                    chunk_topics.append(topic_name)
            
            # Extract decisions relevant to this chunk
            chunk_decisions = []
            for decision in meeting_intelligence.get('decisions_made', []):
                decision_text = decision.get('decision', '')
                if any(word in chunk_content.lower() for word in decision_text.lower().split()[:5]):
                    chunk_decisions.append(decision)
            
            # Extract action items relevant to this chunk
            chunk_actions = []
            for action in meeting_intelligence.get('action_items', []):
                action_text = action.get('task', '')
                if any(word in chunk_content.lower() for word in action_text.lower().split()[:5]):
                    chunk_actions.append(action)
            
            # Extract Q&A relevant to this chunk
            chunk_questions = []
            for qa in meeting_intelligence.get('questions_and_answers', []):
                question_text = qa.get('question', '')
                if any(word in chunk_content.lower() for word in question_text.lower().split()[:5]):
                    chunk_questions.append(qa)
            
            # Determine chunk type based on content
            chunk_type = 'discussion'
            if chunk_decisions:
                chunk_type = 'decision'
            elif chunk_actions:
                chunk_type = 'action_planning'
            elif chunk_questions:
                chunk_type = 'qa'
            
            # Calculate importance score
            importance_score = 0.5
            if chunk_decisions:
                importance_score += 0.3
            if chunk_actions:
                importance_score += 0.2
            if len(chunk_speakers) > 1:
                importance_score += 0.1
            importance_score = min(importance_score, 1.0)
            
            # Extract key phrases from chunk
            chunk_phrases = []
            searchable_metadata = meeting_intelligence.get('searchable_metadata', {})
            for phrase in searchable_metadata.get('key_phrases', []):
                if phrase.lower() in chunk_content.lower():
                    chunk_phrases.append(phrase)
            
            # Set context
            context_before = f"Chunk {chunk_index + 1} of {total_chunks}"
            context_after = f"Part of {meeting_intelligence.get('meeting_metadata', {}).get('main_purpose', 'meeting discussion')}"
            
            # Update chunk intelligence
            chunk_intel.update({
                'chunk_type': chunk_type,
                'speakers': chunk_speakers,
                'speaker_contributions': chunk_speaker_contributions,
                'topics': chunk_topics,
                'decisions': chunk_decisions,
                'actions': chunk_actions,
                'questions': chunk_questions,
                'context_before': context_before,
                'context_after': context_after,
                'key_phrases': chunk_phrases,
                'importance_score': importance_score
            })
            
        except Exception as e:
            logger.error(f"Error extracting chunk intelligence: {e}")
        
        return chunk_intel
    
    def _create_enhanced_content(self, original_content: str, chunk_intelligence: Dict) -> str:
        """Create enhanced content with context and speaker attribution"""
        enhanced_parts = []
        
        # Add context header
        context_before = chunk_intelligence.get('context_before', '')
        if context_before:
            enhanced_parts.append(f"[Context: {context_before}]")
        
        # Add speaker information
        speakers = chunk_intelligence.get('speakers', [])
        if speakers:
            enhanced_parts.append(f"[Speakers: {', '.join(speakers)}]")
        
        # Add chunk type
        chunk_type = chunk_intelligence.get('chunk_type', 'discussion')
        enhanced_parts.append(f"[Type: {chunk_type}]")
        
        # Add original content
        enhanced_parts.append(original_content)
        
        # Add summary of decisions/actions if present
        decisions = chunk_intelligence.get('decisions', [])
        if decisions:
            decision_summary = "; ".join([d.get('decision', '') for d in decisions[:2]])
            enhanced_parts.append(f"[Decisions: {decision_summary}]")
        
        actions = chunk_intelligence.get('actions', [])
        if actions:
            action_summary = "; ".join([a.get('task', '') for a in actions[:2]])
            enhanced_parts.append(f"[Actions: {action_summary}]")
        
        return "\n".join(enhanced_parts)
    
    def process_documents(self, document_folder: str, user_id: str = "default_user", project_id: str = None) -> Dict[str, Any]:
        """Process all documents in a folder with chunking and vector storage
        
        Args:
            document_folder: Path to folder containing documents
            user_id: User ID for document ownership (defaults to 'default_user')
            project_id: Optional project ID for document categorization
            
        Returns:
            Dictionary with processing results and statistics
        """
        folder_path = Path(document_folder)
        
        if not folder_path.exists():
            logger.error(f"Folder {document_folder} does not exist")
            return {"success": False, "error": f"Folder {document_folder} does not exist", "processed_count": 0}
        
        # Get supported files
        supported_extensions = ['.docx', '.txt', '.pdf']
        doc_files = []
        for ext in supported_extensions:
            doc_files.extend(folder_path.glob(f"*{ext}"))
        
        
        processed_count = 0
        for doc_file in doc_files:
            try:
                logger.info(f"Processing: {doc_file.name}")
                
                # Read document content
                content = self.read_document_content(str(doc_file))
                if not content.strip():
                    logger.warning(f"No content extracted from {doc_file.name}")
                    continue
                
                # Parse document
                meeting_doc = self.parse_document_content(content, doc_file.name, user_id, project_id)
                
                # Create chunks with embeddings
                logger.info(f"Creating chunks for {doc_file.name}")
                chunks = self.chunk_document(meeting_doc)
                
                # Store in vector database
                self.vector_db.add_document(meeting_doc, chunks)
                
                processed_count += 1
                logger.info(f"Successfully processed: {doc_file.name} ({len(chunks)} chunks)")
                
            except Exception as e:
                logger.error(f"Error processing {doc_file.name}: {e}")
        
        # Save vector index
        if processed_count > 0:
            self.vector_db.save_index()
            logger.info(f"Successfully processed {processed_count} documents")
        
        return {
            "success": processed_count > 0,
            "processed_count": processed_count,
            "total_files": len(doc_files),
            "user_id": user_id,
            "project_id": project_id
        }
    
    def hybrid_search(self, query: str, user_id: str, project_id: str = None, meeting_id: str = None, folder_path: str = None, top_k: int = 15, semantic_weight: float = 0.7) -> List[DocumentChunk]:
        """Perform hybrid search combining semantic and keyword search"""
        
        # Extract keywords from query
        keywords = [word.lower().strip() for word in query.split() if len(word) > 2]
        
        # Semantic search
        try:
            if self.embedding_model is None:
                logger.warning("Embedding model not available, skipping semantic search")
                semantic_results = []
            else:
                query_embedding = np.array(self.embedding_model.embed_query(query))
                if folder_path:
                    # Filter semantic search by folder
                    semantic_results = self.vector_db.search_similar_chunks_by_folder(query_embedding, user_id, folder_path, top_k * 2)
                else:
                    semantic_results = self.vector_db.search_similar_chunks(query_embedding, top_k * 2)
        except Exception as e:
            logger.error(f"Error in semantic search: {e}")
            semantic_results = []
        
        # Keyword search with user context
        if folder_path:
            keyword_chunk_ids = self.vector_db.keyword_search_chunks_by_folder(keywords, user_id, folder_path, top_k)
        else:
            keyword_chunk_ids = self.vector_db.keyword_search_chunks_by_user(keywords, user_id, project_id, meeting_id, top_k)
        
        # Combine and score results
        chunk_scores = defaultdict(float)
        
        # Add semantic scores
        for chunk_id, similarity in semantic_results:
            chunk_scores[chunk_id] += similarity * semantic_weight
        
        # Add keyword scores
        keyword_weight = 1.0 - semantic_weight
        for chunk_id in keyword_chunk_ids:
            chunk_scores[chunk_id] += keyword_weight * 0.5  # Base keyword score
        
        # Get top chunks
        top_chunk_ids = sorted(chunk_scores.keys(), key=lambda x: chunk_scores[x], reverse=True)[:top_k]
        
        # Retrieve chunk details
        return self.vector_db.get_chunks_by_ids(top_chunk_ids)
    
    def detect_summary_query(self, query: str) -> bool:
        """Detect if the query is asking for meeting summaries"""
        summary_keywords = [
            'summarize', 'summary', 'summaries', 'overview', 'brief', 
            'recap', 'highlights', 'key points', 'main points',
            'all meetings', 'all documents', 'overall', 'across all',
            'consolidate', 'aggregate', 'compile', 'comprehensive',
            'meetings summary', 'meeting summaries', 'summarize meetings',
            'summarize the meetings', 'summary of meetings', 'summary of all'
        ]
        
        query_lower = query.lower()
        for keyword in summary_keywords:
            if keyword in query_lower:
                return True
        return False

    def answer_query_with_intelligence(self, query: str, user_id: str, document_ids: List[str] = None, 
                                     project_id: str = None, meeting_id: str = None, meeting_ids: List[str] = None, 
                                     date_filters: Dict[str, Any] = None, folder_path: str = None, 
                                     context_limit: int = 10, include_context: bool = False) -> Union[str, Tuple[str, str]]:
        """Answer user query using enhanced intelligence-aware search and context reconstruction"""
        
        try:
            # ===== DEBUG LOGGING: MAIN PROCESSOR ENTRY =====
            logger.info("[PROCESSOR] MeetingProcessor.answer_query_with_intelligence() - ENTRY POINT")
            logger.info("=" * 80)
            logger.info(f"[PROCESSING] QUERY: '{query}'")
            logger.info(f"[USER] User ID: {user_id}")
            logger.info(f"[PARAMS] Parameters:")
            logger.info(f"   - document_ids: {document_ids}")
            logger.info(f"   - project_id: {project_id}")
            logger.info(f"   - meeting_id: {meeting_id}")
            logger.info(f"   - meeting_ids: {meeting_ids}")
            logger.info(f"   - date_filters: {date_filters}")
            logger.info(f"   - folder_path: {folder_path}")
            logger.info(f"   - context_limit: {context_limit}")
            logger.info("=" * 80)
            
            # Generate query embedding
            logger.info("[STEPA] Checking embedding model availability...")
            if self.embedding_model is None:
                logger.error("[CRITICAL] Embedding model not available")
                return "Sorry, the system is not properly configured for queries.", ""
            
            logger.info("[OK] Embedding model available - generating query embedding...")
            query_embedding = self.embedding_model.embed_query(query)
            query_vector = np.array(query_embedding)
            logger.info(f"[EMBEDDING] Query embedding generated: {query_vector.shape} dimensions")
            
            # Analyze query to determine filters
            logger.info("[STEPB] Analyzing query for automatic filters...")
            search_filters = self._analyze_query_for_filters(query)
            logger.info(f"[AUTO] Automatic filters detected: {search_filters}")
            
            # Apply user context filters
            logger.info("[STEPC] Applying user-provided filters...")
            original_filters = dict(search_filters)  # Keep copy for comparison
            
            if meeting_ids:
                search_filters['meeting_ids'] = meeting_ids
            if project_id:
                search_filters['project_id'] = project_id
            if date_filters:
                search_filters['date_filters'] = date_filters
            if folder_path:
                search_filters['folder_path'] = folder_path
            
            logger.info(f"[FILTERS] Filter combination:")
            logger.info(f"   - Automatic filters: {original_filters}")
            logger.info(f"   - Final filters: {search_filters}")
            
            # Perform enhanced search with metadata filtering
            logger.info("[STEPD] Starting ENHANCED SEARCH with metadata filtering...")
            logger.info(f"   -> Calling vector_db.enhanced_search_with_metadata()")
            logger.info(f"   -> Query vector shape: {query_vector.shape}")
            logger.info(f"   -> User ID: {user_id}")
            logger.info(f"   -> Filters: {search_filters}")
            logger.info(f"   -> Top K: {context_limit}")
            
            enhanced_results = self.vector_db.enhanced_search_with_metadata(
                query_vector, user_id, search_filters, top_k=context_limit
            )
            
            # ===== DEBUG LOGGING: ENHANCED SEARCH RESULTS =====
            logger.info("[RESULTS] ENHANCED SEARCH RESULTS ANALYSIS")
            logger.info(f"[COUNT] Enhanced search returned: {len(enhanced_results) if enhanced_results else 0} results")
            if enhanced_results:
                logger.info(f"[SAMPLE] Sample results (first 3):")
                for i, result in enumerate(enhanced_results[:3]):
                    if isinstance(result, dict):
                        chunk = result.get('chunk', {})
                        score = result.get('similarity_score', 0.0)
                        chunk_id = chunk.get('chunk_id', 'unknown') if isinstance(chunk, dict) else 'unknown'
                        logger.info(f"   {i+1}. Chunk ID: {chunk_id}, Score: {score:.4f}")
                    else:
                        logger.info(f"   {i+1}. Result: {type(result)} - {str(result)[:100]}...")
            else:
                logger.error("[ERROR] ENHANCED SEARCH RETURNED ZERO RESULTS!")
                logger.error("   -> This is the root cause of 'no relevant information' responses!")
            
            logger.info(f"[APPLIED] Search filters that were applied: {search_filters}")
            
            if not enhanced_results:
                logger.warning("[FALLBACK] TRIGGERED: Enhanced search returned no results")
                logger.warning("   -> Starting BASIC SEARCH without metadata filters...")
                
                # Fallback to basic search without metadata filters
                try:
                    logger.info("[STEP E: BASIC SEARCH FALLBACK")
                    logger.info("   -> Calling vector_db.search_similar_chunks() (no filters)")
                    
                    basic_results = self.vector_db.search_similar_chunks(query_vector, top_k=context_limit)
                    
                    logger.info(f"[STATS] Basic search returned: {len(basic_results) if basic_results else 0} raw results")
                    
                    if basic_results:
                        logger.info("[STEP F: Converting basic results to chunks...")
                        chunk_ids = [chunk_id for chunk_id, _ in basic_results]
                        logger.info(f"Chunk IDs from basic search: {chunk_ids[:3]}... (showing first 3)")
                        
                        chunks = self.vector_db.get_chunks_by_ids(chunk_ids)
                        
                        logger.info(f"Retrieved {len(chunks) if chunks else 0} chunks from database")
                        
                        if chunks:
                            logger.info("[OK] Basic search fallback successful - converting to enhanced format")
                            
                            # Convert to enhanced format for compatibility
                            score_map = {chunk_id: score for chunk_id, score in basic_results}
                            enhanced_format = []
                            
                            for i, chunk in enumerate(chunks):
                                # Handle case where chunk might be a string or dict instead of object
                                if isinstance(chunk, str):
                                    logger.error(f"[ERROR] Unexpected string chunk in fallback: {chunk}")
                                    continue
                                elif isinstance(chunk, dict):
                                    chunk_id = chunk.get('chunk_id', '')
                                    enhanced_format.append({
                                        'chunk': chunk,
                                        'similarity_score': score_map.get(chunk_id, 0.0),
                                        'context': ''
                                    })
                                    if i < 2:  # Log first 2 chunks
                                        logger.info(f"   Chunk {i+1}: ID={chunk_id}, Score={score_map.get(chunk_id, 0.0):.4f}")
                                else:
                                    # Assume it's a proper chunk object
                                    chunk_id = getattr(chunk, 'chunk_id', 'unknown')
                                    enhanced_format.append({
                                        'chunk': chunk,
                                        'similarity_score': score_map.get(chunk_id, 0.0),
                                        'context': ''
                                    })
                            
                            logger.info(f"[OK] Converted {len(enhanced_format)} chunks to enhanced format for fallback")
                            logger.info("[STEP G: Generating response from fallback results...")
                            
                            response, context = self._generate_intelligence_response(query, enhanced_format, user_id)
                            
                            logger.info("[RESULT] FALLBACK RESPONSE GENERATED SUCCESSFULLY")
                            return (response, context) if include_context else response
                        else:
                            logger.error("[ERROR] Basic search returned chunk IDs but couldn't retrieve chunk data")
                    else:
                        logger.error("[ERROR] Basic search also returned zero results")
                        
                except Exception as e:
                    logger.error(f"[ERROR] FALLBACK SEARCH FAILED: {e}")
                    import traceback
                    logger.error(f"Traceback: {traceback.format_exc()}")
                
                logger.error("[ALERT] BOTH ENHANCED AND BASIC SEARCH FAILED - RETURNING NO INFORMATION RESPONSE")
                return "I couldn't find any relevant information for your query. Please try rephrasing or check if you have uploaded meeting documents.", ""
            
            # Generate context-aware response
            logger.info("[STEP H: ENHANCED SEARCH SUCCESSFUL - Generating response...")
            logger.info(f"   -> Calling _generate_intelligence_response() with {len(enhanced_results)} results")
            
            response, context = self._generate_intelligence_response(query, enhanced_results, user_id)
            
            logger.info("[RESULT] ENHANCED SEARCH RESPONSE GENERATED SUCCESSFULLY")
            logger.info(f"Final response length: {len(response)} characters")
            logger.info(f"Final context length: {len(context) if context else 0} characters")
            
            if include_context:
                return response, context
            else:
                return response
                
        except Exception as e:
            logger.error(f"Error in intelligence-aware query processing: {e}")
            return f"I encountered an error while processing your query: {str(e)}", ""
    
    def _analyze_query_for_filters(self, query: str) -> Dict:
        """Analyze query to determine appropriate metadata filters"""
        filters = {}
        query_lower = query.lower()
        
        # Detect speaker-specific queries
        speaker_patterns = ['what did', 'what said', 'who said', 'mentioned by', 'according to']
        if any(pattern in query_lower for pattern in speaker_patterns):
            # Try to extract speaker names from query
            potential_speakers = []
            words = query.split()
            
            # Enhanced pattern matching for speaker extraction
            for i, word in enumerate(words):
                word_lower = word.lower()
                if word_lower in ['what', 'who'] and i + 1 < len(words):
                    next_word = words[i + 1].lower()
                    if next_word in ['did', 'said']:
                        # Look for the speaker name after "what did" or "who said"
                        if i + 2 < len(words):
                            speaker_name = words[i + 2].strip('.,?!').title()
                            potential_speakers.append(speaker_name)
                
                # Also check for patterns like "sandeep said" or "according to john"
                if word_lower in ['said', 'mentioned'] and i > 0:
                    speaker_name = words[i - 1].strip('.,?!').title()
                    potential_speakers.append(speaker_name)
                elif word_lower == 'to' and i > 0 and words[i - 1].lower() == 'according':
                    if i + 1 < len(words):
                        speaker_name = words[i + 1].strip('.,?!').title()
                        potential_speakers.append(speaker_name)
            
            # Remove duplicates and common words
            common_words = {'the', 'in', 'at', 'on', 'for', 'with', 'by', 'from', 'meeting', 'discussion'}
            potential_speakers = list(set([s for s in potential_speakers if s.lower() not in common_words]))
            
            if potential_speakers:
                filters['speakers'] = potential_speakers
                logger.info(f"Extracted speakers from query '{query}': {potential_speakers}")
        
        # Detect decision-focused queries
        decision_patterns = ['decision', 'decided', 'conclusion', 'resolution', 'agreed']
        if any(pattern in query_lower for pattern in decision_patterns):
            filters['has_decisions'] = True
            filters['chunk_type'] = ['decision', 'discussion']
        
        # Detect action item queries
        action_patterns = ['action', 'task', 'todo', 'follow up', 'next steps', 'assigned']
        if any(pattern in query_lower for pattern in action_patterns):
            filters['has_actions'] = True
            filters['chunk_type'] = ['action_planning', 'discussion']
        
        # Detect high-importance queries
        importance_patterns = ['important', 'critical', 'urgent', 'key', 'major']
        if any(pattern in query_lower for pattern in importance_patterns):
            filters['min_importance'] = 0.7
        
        return filters
    
    def _analyze_context_richness(self, enhanced_results: List[Dict]) -> Dict[str, Any]:
        """Analyze the richness and quality of available context"""
        speakers = set()
        meetings = set()
        total_content_length = 0
        decisions_count = 0
        actions_count = 0
        time_span_days = 0
        dates = []
        
        for result in enhanced_results:
            chunk = result['chunk']
            
            # Collect content statistics
            total_content_length += len(chunk.content) if chunk.content else 0
            
            # Extract speakers
            if hasattr(chunk, 'speakers') and chunk.speakers:
                try:
                    import json
                    chunk_speakers = json.loads(chunk.speakers) if isinstance(chunk.speakers, str) else chunk.speakers
                    speakers.update(chunk_speakers)
                except:
                    pass
            
            # Extract meetings and dates
            if hasattr(chunk, 'filename') and chunk.filename:
                meetings.add(chunk.filename)
            
            if hasattr(chunk, 'date') and chunk.date:
                dates.append(chunk.date)
            
            # Count decisions and actions
            if hasattr(chunk, 'decisions') and chunk.decisions:
                try:
                    import json
                    decisions = json.loads(chunk.decisions) if isinstance(chunk.decisions, str) else chunk.decisions
                    decisions_count += len(decisions) if decisions else 0
                except:
                    pass
            
            if hasattr(chunk, 'actions') and chunk.actions:
                try:
                    import json
                    actions = json.loads(chunk.actions) if isinstance(chunk.actions, str) else chunk.actions
                    actions_count += len(actions) if actions else 0
                except:
                    pass
        
        # Calculate time span
        if dates and len(dates) > 1:
            try:
                sorted_dates = sorted([d for d in dates if d])
                if len(sorted_dates) >= 2:
                    time_span_days = (sorted_dates[-1] - sorted_dates[0]).days
            except:
                time_span_days = 0
        
        return {
            'speakers_count': len(speakers),
            'meetings_count': len(meetings),
            'total_content_length': total_content_length,
            'decisions_count': decisions_count,
            'actions_count': actions_count,
            'time_span_days': time_span_days,
            'context_richness_score': min(100, (len(speakers) * 10) + (len(meetings) * 5) + (total_content_length // 100))
        }

    def _determine_response_requirements(self, query: str, context_analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Determine dynamic response requirements based on query and context"""
        query_lower = query.lower()
        
        # Base requirements
        min_words = 200
        max_words = 800
        detail_level = "comprehensive"
        
        # Adjust based on context richness
        richness_score = context_analysis['context_richness_score']
        if richness_score > 75:
            min_words = 400
            max_words = 1000
            detail_level = "extensive"
        elif richness_score > 40:
            min_words = 300
            max_words = 600
            detail_level = "detailed"
        
        # Adjust based on query complexity
        complex_indicators = ['explain', 'how', 'why', 'process', 'background', 'context', 'detailed', 'comprehensive']
        technical_indicators = ['technical', 'implementation', 'architecture', 'design', 'algorithm', 'method']
        timeline_indicators = ['timeline', 'sequence', 'chronological', 'history', 'progression', 'development']
        
        if any(indicator in query_lower for indicator in complex_indicators):
            min_words += 100
            detail_level = "comprehensive"
        
        if any(indicator in query_lower for indicator in technical_indicators):
            min_words += 150
            detail_level = "technical_detailed"
        
        if any(indicator in query_lower for indicator in timeline_indicators):
            min_words += 100
            detail_level = "chronological_detailed"
        
        return {
            'min_words': min_words,
            'max_words': max_words,
            'detail_level': detail_level,
            'context_richness': richness_score
        }

    def _generate_intelligence_response(self, query: str, enhanced_results: List[Dict], user_id: str) -> Tuple[str, str]:
        """Generate comprehensive response using enhanced intelligence data with dynamic scaling"""
        
        # Analyze context richness and determine response requirements
        context_analysis = self._analyze_context_richness(enhanced_results)
        response_requirements = self._determine_response_requirements(query, context_analysis)
        
        logger.info(f"Context analysis: {context_analysis}")
        logger.info(f"Response requirements: {response_requirements}")
        
        # Prepare comprehensive context from enhanced results
        context_parts = []
        speaker_contributions = {}
        decisions_made = []
        actions_identified = []
        meeting_contexts = []
        
        for result in enhanced_results:
            chunk = result['chunk']
            context = result['context']
            # Extract metadata from chunk attributes instead of result['metadata'] which doesn't exist
            metadata = {
                'speakers': [],
                'decisions': [],
                'actions': []
            }
            
            # Try to extract metadata from chunk attributes if available
            if hasattr(chunk, 'speakers') and chunk.speakers:
                try:
                    import json
                    metadata['speakers'] = json.loads(chunk.speakers) if isinstance(chunk.speakers, str) else chunk.speakers
                except:
                    metadata['speakers'] = []
            
            if hasattr(chunk, 'decisions') and chunk.decisions:
                try:
                    import json
                    metadata['decisions'] = json.loads(chunk.decisions) if isinstance(chunk.decisions, str) else chunk.decisions
                except:
                    metadata['decisions'] = []
                    
            if hasattr(chunk, 'actions') and chunk.actions:
                try:
                    import json
                    metadata['actions'] = json.loads(chunk.actions) if isinstance(chunk.actions, str) else chunk.actions
                except:
                    metadata['actions'] = []
            
            # Collect context information
            document_title = context.get('document_title', chunk.filename)
            document_date = context.get('document_date', 'Unknown date')
            
            context_part = f"**From {document_title} ({document_date})**\n"
            context_part += f"Content: {chunk.enhanced_content or chunk.content}\n"
            
            if metadata['speakers']:
                context_part += f"Speakers: {', '.join(metadata['speakers'])}\n"
            
            if metadata['decisions']:
                context_part += f"Decisions: {'; '.join([str(d) for d in metadata['decisions'][:2]])}\n"
                decisions_made.extend(metadata['decisions'])
            
            if metadata['actions']:
                context_part += f"Actions: {'; '.join([str(a) for a in metadata['actions'][:2]])}\n"
                actions_identified.extend(metadata['actions'])
            
            context_part += f"Relevance Score: {result['similarity_score']:.3f}\n\n"
            context_parts.append(context_part)
            
            # Collect speaker data
            for speaker in metadata['speakers']:
                if speaker not in speaker_contributions:
                    speaker_contributions[speaker] = []
                speaker_contributions[speaker].append({
                    'content': chunk.content[:200] + "..." if len(chunk.content) > 200 else chunk.content,
                    'meeting': document_title,
                    'date': document_date
                })
            
            # Collect meeting context
            meeting_context = {
                'meeting_title': document_title,
                'meeting_date': document_date,
                'participants': context.get('participants', ''),
                'main_topics': context.get('main_topics', ''),
                'document_summary': context.get('document_summary', '')
            }
            if meeting_context not in meeting_contexts:
                meeting_contexts.append(meeting_context)
        
        # Create comprehensive context string
        full_context = "COMPREHENSIVE MEETING INTELLIGENCE:\n\n" + "\n".join(context_parts)
        
        # Generate enhanced response prompt
        response_prompt = f"""
        Based on the comprehensive meeting intelligence provided, answer the user's query with complete accuracy and context.
        
        USER QUERY: "{query}"
        
        MEETING INTELLIGENCE CONTEXT:
        {full_context}
        
        SPEAKER ANALYSIS:
        {self._format_speaker_analysis(speaker_contributions)}
        
        DECISIONS SUMMARY:
        {self._format_decisions_summary(decisions_made)}
        
        ACTION ITEMS SUMMARY:
        {self._format_actions_summary(actions_identified)}
        
        INSTRUCTIONS:
        1. Answer the user's query directly and comprehensively
        2. Include specific speaker attributions when relevant
        3. Reference specific meetings and dates
        4. Include related decisions and action items
        5. Maintain chronological context when applicable
        6. Use the enhanced context to provide complete information
        
        Provide a detailed, accurate response that leverages all the intelligence available.
        """
        
        try:
            messages = [
                SystemMessage(content="You are an expert meeting intelligence assistant. Provide comprehensive, accurate responses using all available context and speaker attributions."),
                HumanMessage(content=response_prompt)
            ]
            
            response = self.llm.invoke(messages)
            response_content = response.content.strip()
            logger.info(f"LLM response generated successfully. Length: {len(response_content)} characters")
            logger.info(f"Response preview: {response_content[:200]}...")
            return response_content, full_context
            
        except Exception as e:
            logger.error(f"Error generating intelligence response: {e}")
            return f"I found relevant information but encountered an error generating the response: {str(e)}", full_context
    
    def _format_speaker_analysis(self, speaker_contributions: Dict) -> str:
        """Format speaker contribution analysis"""
        if not speaker_contributions:
            return "No specific speaker contributions identified."
        
        analysis = []
        for speaker, contributions in speaker_contributions.items():
            contrib_summary = f"{speaker}: {len(contributions)} contributions across {len(set(c['meeting'] for c in contributions))} meetings"
            analysis.append(contrib_summary)
        
        return "\n".join(analysis)
    
    def _format_decisions_summary(self, decisions: List[Dict]) -> str:
        """Format decisions summary"""
        if not decisions:
            return "No specific decisions identified."
        
        summary = []
        for decision in decisions[:5]:  # Limit to top 5
            dec_text = decision.get('decision', 'Unknown decision')
            decided_by = decision.get('decided_by', 'Unknown')
            summary.append(f"- {dec_text} (by {decided_by})")
        
        return "\n".join(summary)
    
    def _format_actions_summary(self, actions: List[Dict]) -> str:
        """Format action items summary"""
        if not actions:
            return "No specific action items identified."
        
        summary = []
        for action in actions[:5]:  # Limit to top 5
            task = action.get('task', 'Unknown task')
            assigned_to = action.get('assigned_to', 'Unknown')
            summary.append(f"- {task} (assigned to {assigned_to})")
        
        return "\n".join(summary)

    def answer_query(self, query: str, user_id: str, document_ids: List[str] = None, project_id: str = None, meeting_id: str = None, meeting_ids: List[str] = None, date_filters: List[str] = None, folder_path: str = None, context_limit: int = 10, include_context: bool = False) -> Union[str, Tuple[str, str]]:
        """Answer user query using hybrid search and intelligent context selection"""
        
        # Check if this is a summary query and no specific documents are selected
        is_summary_query = self.detect_summary_query(query)
        
        # Check if this is a comprehensive project summary query
        is_project_summary_query = self.detect_project_summary_query(query)
        
        # Handle comprehensive project summary requests
        if is_project_summary_query and not document_ids:
            logger.info("Detected comprehensive project summary query")
            return self._generate_comprehensive_project_summary(query, user_id, project_id, include_context)
        
        # Handle enhanced @ mention filters
        if meeting_ids:
            logger.info(f"Enhanced meeting filters: {meeting_ids}")
            # Support for multiple meeting IDs - use all provided meeting IDs
            if meeting_ids and not meeting_id:
                # Use all meeting IDs for document retrieval
                meeting_id = meeting_ids  # Pass the full list instead of just first one
        
        if date_filters:
            logger.info(f"Date filters: {date_filters}")
            # Apply explicit date filtering using existing timeframe logic
            date_filtered_docs = []
            for date_filter in date_filters:
                try:
                    # Use existing timeframe detection and date filtering
                    timeframe_docs = self.get_documents_by_timeframe(date_filter, user_id)
                    # Extract document IDs from MeetingDocument objects
                    filtered_doc_ids = [doc.document_id for doc in timeframe_docs]
                    date_filtered_docs.extend(filtered_doc_ids)
                    logger.info(f"Date filter '{date_filter}' matched {len(filtered_doc_ids)} documents")
                except Exception as e:
                    logger.warning(f"Error applying date filter '{date_filter}': {e}")
            
            # Remove duplicates and use date-filtered documents
            if date_filtered_docs:
                date_filtered_docs = list(set(date_filtered_docs))
                if not document_ids:
                    document_ids = date_filtered_docs
                else:
                    # Intersect with existing document filters
                    document_ids = list(set(document_ids) & set(date_filtered_docs))
                logger.info(f"Applied date filters, now using {len(document_ids)} documents")
        
        # If no specific documents are provided, include all user documents
        if not document_ids:
            logger.info("No specific documents provided, including all user documents")
            if folder_path:
                # Use folder-based filtering ONLY - don't mix with project/meeting filters
                document_ids = self.vector_db.get_user_documents_by_folder(user_id, folder_path)
            else:
                # Use standard scope-based filtering
                document_ids = self.vector_db.get_user_documents_by_scope(user_id, project_id, meeting_id)
                logger.info(f"Including {len(document_ids)} documents for user {user_id}")
        
        # Enhanced intelligent timeframe detection
        detected_timeframe = self._detect_timeframe_from_query(query)
        
        # Get relevant documents by timeframe if specified
        if detected_timeframe:
            logger.info(f"Detected timeframe: {detected_timeframe}")
            timeframe_docs = self.vector_db.get_documents_by_timeframe(detected_timeframe, user_id)
            if not timeframe_docs:
                error_msg = f"I don't have any meeting documents from the {detected_timeframe.replace('_', ' ')} timeframe."
                return (error_msg, "") if include_context else error_msg
            
            # Filter document_ids to only include documents from the detected timeframe
            if not document_ids:
                document_ids = [doc.document_id for doc in timeframe_docs]
                logger.info(f"Using {len(document_ids)} documents from {detected_timeframe} timeframe")
            else:
                # Intersect with timeframe documents
                timeframe_doc_ids = {doc.document_id for doc in timeframe_docs}
                document_ids = [doc_id for doc_id in document_ids if doc_id in timeframe_doc_ids]
                logger.info(f"Filtered to {len(document_ids)} documents in {detected_timeframe} timeframe")
            
            # Generate enhanced summary if this is a summary query with date context
            if is_summary_query and len(timeframe_docs) > 1:
                return self._generate_date_based_summary(query, timeframe_docs, detected_timeframe, include_context)
        
        # Perform hybrid search
        relevant_chunks = self.hybrid_search(query, user_id, project_id, meeting_id, folder_path, top_k=context_limit * 3)
        
        # Filter chunks by document IDs if specified
        if document_ids:  # Temporarily disabled @ feature
            original_count = len(relevant_chunks)
            relevant_chunks = [chunk for chunk in relevant_chunks if chunk.document_id in document_ids]
            if not relevant_chunks:
                error_msg = "I don't have any relevant information in the specified documents for your question."
                return (error_msg, "") if include_context else error_msg
        elif not relevant_chunks:
            error_msg = "I don't have any relevant meeting documents to answer your question."
            return (error_msg, "") if include_context else error_msg
        
        # Group chunks by document and select best representatives
        document_chunks = defaultdict(list)
        for chunk in relevant_chunks:
            document_chunks[chunk.document_id].append(chunk)
        
        # Select best chunks from each document (more for summary queries)
        selected_chunks = []
        chunks_per_doc = 5 if is_summary_query else 2  # More chunks for summaries
        for doc_id, chunks in document_chunks.items():
            # Sort chunks by position to maintain context
            chunks.sort(key=lambda x: x.chunk_index)
            selected_chunks.extend(chunks[:chunks_per_doc])  # More chunks for summaries
        
        # Limit total chunks
        selected_chunks = selected_chunks[:context_limit]
        
        # Build context without chunk references
        context_parts = []
        current_doc = None
        
        for chunk in selected_chunks:
            if chunk.document_id != current_doc:
                # Add document header when switching documents
                if current_doc is not None:
                    context_parts.append("\n" + "="*60 + "\n")
                
                context_parts.append(f"Document: {chunk.filename}")
                current_doc = chunk.document_id
            
            # Add content without chunk numbering
            context_parts.append(chunk.content)
        
        context = "\n".join(context_parts)
        
        # Generate enhanced prompt for summary queries
        if is_summary_query:
            answer_prompt = f"""
User Question: {query}

Meeting Document Context:
{context}

This is a summary request. Please provide a comprehensive overview that includes:
- Key topics and themes discussed across meetings
- Important decisions made and their context
- Action items and next steps identified
- Any patterns or trends across the meetings
- Significant outcomes or conclusions

Organize the information in a clear, detailed manner. Be thorough - the user wants a complete picture, not just highlights.
"""
        else:
            answer_prompt = f"""
User Question: {query}

Meeting Document Context:
{context}

Please provide a detailed, conversational response to the user's question based on the meeting documents. Be thorough and comprehensive in your answer, including specific details, quotes, and context from the meetings. Avoid using structured formats with bullet points or section headers like "Key Topics" or "Decisions" unless the user specifically asks for that format.

Write your response as if you're having a natural conversation with the user, providing rich detail and specific information from the meetings. Include relevant background context, specific quotes or examples, and elaborate on the important points.

IMPORTANT: When referencing information from the documents, always cite the document filename (e.g., "Document_Fulfillment_AIML-20250714_153021-Meeting_Recording.docx") rather than chunk numbers. This helps users know which specific document the information comes from.
"""
        
        try:
            messages = [
                SystemMessage(content="You are a helpful AI assistant that provides detailed, conversational responses about meeting documents. Avoid structured formats with bullet points or headers unless specifically requested. Provide rich, comprehensive answers with specific details, quotes, and context. Always cite document filenames rather than chunk numbers when referencing information."),
                HumanMessage(content=answer_prompt)
            ]
            
            response = self.llm.invoke(messages)
            response_content = response.content.strip()
            logger.info(f"Standard LLM response generated. Length: {len(response_content)} characters")
            logger.info(f"Response preview: {response_content[:200]}...")
            return (response_content, context) if include_context else response_content
            
        except Exception as e:
            logger.error(f"Error generating answer: {e}")
            # Try refreshing clients and retry
            try:
                self.refresh_clients()
                messages = [
                    SystemMessage(content="You are a helpful AI assistant that provides detailed, conversational responses about meeting documents. Avoid structured formats with bullet points or headers unless specifically requested. Provide rich, comprehensive answers with specific details, quotes, and context. Always cite document filenames rather than chunk numbers when referencing information."),
                    HumanMessage(content=answer_prompt)
                ]
                response = self.llm.invoke(messages)
                return (response.content, context) if include_context else response.content
            except Exception as retry_error:
                error_msg = f"I encountered an error while processing your question. Please try again later. Error details: {str(retry_error)}"
                return (error_msg, "") if include_context else error_msg
    
    def generate_follow_up_questions(self, user_query: str, ai_response: str, context: str) -> List[str]:
        """Generate follow-up questions based on the user query, AI response, and document context"""
        try:
            follow_up_prompt = f"""
Based on the user's question, the AI response provided, and the meeting document context, generate 4-5 relevant follow-up questions that the user might want to ask next.

User's Question: {user_query}

AI Response: {ai_response}

Meeting Context: {context[:2000]}...

Generate follow-up questions that:
- Build upon the current conversation topic
- Explore related aspects mentioned in the documents
- Ask for deeper details about key points
- Inquire about timelines, next steps, or implications
- Connect to other relevant topics in the meetings

Return exactly 4-5 questions, each on a new line, without numbers or bullet points. Make them natural and conversational.
"""

            messages = [
                SystemMessage(content="You generate natural, conversational follow-up questions that help users explore their meeting documents further."),
                HumanMessage(content=follow_up_prompt)
            ]
            
            response = self.llm.invoke(messages)
            
            # Parse the response into individual questions
            questions = [q.strip() for q in response.content.split('\n') if q.strip() and len(q.strip()) > 10]
            
            # Ensure we have 4-5 questions
            if len(questions) > 5:
                questions = questions[:5]
            elif len(questions) < 4:
                # Add generic questions if we don't have enough
                generic_questions = [
                    "What were the key decisions made in these meetings?",
                    "What are the next steps mentioned?",
                    "Are there any deadlines or milestones discussed?"
                ]
                questions.extend(generic_questions[:5-len(questions)])
            
            return questions[:5]
            
        except Exception as e:
            logger.error(f"Error generating follow-up questions: {e}")
            # Return default follow-up questions
            return [
                "What were the main decisions made in the meetings?",
                "What are the upcoming deadlines or milestones?",
                "Are there any action items assigned?"
            ]
    
    def get_meeting_statistics(self) -> Dict[str, Any]:
        """Get simplified statistics about processed meetings (excluding soft-deleted)"""
        try:
            conn = sqlite3.connect(self.vector_db.db_path)
            cursor = conn.cursor()
            
            # Get ACTIVE document counts only (exclude soft-deleted)
            cursor.execute("""
                SELECT COUNT(*) FROM documents 
                WHERE is_deleted = FALSE OR is_deleted IS NULL
            """)
            total_docs = cursor.fetchone()[0]
            
            if total_docs == 0:
                return {"error": "No documents processed"}
            
            # Get date range from ACTIVE meeting dates in filenames
            cursor.execute("""
                SELECT filename FROM documents 
                WHERE is_deleted = FALSE OR is_deleted IS NULL
            """)
            filenames = cursor.fetchall()
            
            meeting_dates = []
            for (filename,) in filenames:
                # Extract date from filename format: Document_Fulfillment_AIML-20250714_153021-Meeting_Recording.docx
                date_match = re.search(r'-(\d{8})_', filename)
                if date_match:
                    date_str = date_match.group(1)  # e.g., "20250714"
                    # Convert to readable format: YYYY-MM-DD
                    formatted_date = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
                    meeting_dates.append(formatted_date)
            
            if meeting_dates:
                meeting_dates.sort()
                date_range = (meeting_dates[0], meeting_dates[-1])
            else:
                # Fallback to processing dates if no meeting dates found (active only)
                cursor.execute("""
                    SELECT MIN(date), MAX(date) FROM documents 
                    WHERE is_deleted = FALSE OR is_deleted IS NULL
                """)
                date_range = cursor.fetchone()
            
            # Get ACTIVE chunk statistics only (exclude soft-deleted)
            cursor.execute("""
                SELECT COUNT(*), AVG(LENGTH(content)) FROM chunks 
                WHERE is_deleted = FALSE OR is_deleted IS NULL
            """)
            chunk_stats = cursor.fetchone()
            
            # Get monthly distribution (optional - can keep for internal tracking)
            cursor.execute("SELECT strftime('%Y-%m', date) as month, COUNT(*) FROM documents GROUP BY month ORDER BY month")
            monthly_counts = dict(cursor.fetchall())
            
            conn.close()
            
            stats = {
                "total_meetings": total_docs,
                "total_chunks": chunk_stats[0] if chunk_stats[0] else 0,
                "average_chunk_length": int(chunk_stats[1]) if chunk_stats[1] else 0,
                "vector_index_size": self.vector_db.index.ntotal if self.vector_db.index else 0,
                "date_range": {
                    "earliest": date_range[0] if date_range[0] else "N/A",
                    "latest": date_range[1] if date_range[1] else "N/A"
                },
                "meetings_per_month": monthly_counts
            }
            
            return stats
        except Exception as e:
            logger.error(f"Error generating statistics: {e}")
            return {"error": f"Failed to generate statistics: {e}"}

    def detect_project_summary_query(self, query: str) -> bool:
        """Detect if the query is asking for a comprehensive project summary"""
        project_summary_keywords = [
            'project summary', 'project summaries', 'summarize project', 'summarize the project',
            'summary of project', 'summary of all files', 'all files summary', 'comprehensive summary',
            'summarize all meetings', 'all meetings summary', 'overall project', 'entire project',
            'project overview', 'complete summary', 'full summary', 'all documents summary',
            'project recap', 'project highlights', 'all files in project', 'everything in project'
        ]
        
        query_lower = query.lower()
        for keyword in project_summary_keywords:
            if keyword in query_lower:
                return True
        return False

    def _generate_comprehensive_project_summary(self, query: str, user_id: str, project_id: str = None, include_context: bool = False) -> Union[str, Tuple[str, str]]:
        """Generate flexible comprehensive answer that processes ALL files based on user query"""
        try:
            logger.info(f"Generating user-centric comprehensive answer for user {user_id}, project {project_id}")
            
            # Get all documents in the project
            if project_id:
                documents = self.vector_db.get_project_documents(project_id, user_id)
            else:
                documents = self.vector_db.get_all_documents(user_id)
            
            if not documents:
                error_msg = "No documents found in the project to analyze."
                return (error_msg, "") if include_context else error_msg
            
            total_files = len(documents)
            logger.info(f"Found {total_files} files to process for user query: '{query}'")
            
            # Use single flexible processing approach
            return self._generate_flexible_comprehensive_answer(documents, query, total_files, include_context)
                
        except Exception as e:
            logger.error(f"Error generating comprehensive answer: {e}")
            error_msg = f"I encountered an error processing your question across all project files: {str(e)}"
            return (error_msg, "") if include_context else error_msg

    def _generate_flexible_comprehensive_answer(self, documents: List[Any], query: str, total_files: int, include_context: bool) -> Union[str, Tuple[str, str]]:
        """Single flexible function that processes ALL files based on user query intent"""
        try:
            logger.info(f"Processing {total_files} files with flexible approach for query: '{query}'")
            
            # Smart content selection based on file count
            if total_files <= 20:
                # Small projects: Use detailed content from all files
                content_chunks = self._get_detailed_content_from_all_files(documents)
                processing_note = f"Analyzed all {total_files} files individually"
            elif total_files <= 50:
                # Medium projects: Use smart sampling + summaries
                content_chunks = self._get_smart_sampled_content(documents, query)
                processing_note = f"Analyzed all {total_files} files using smart sampling"
            else:
                # Large projects: Use summarized content + key excerpts
                content_chunks = self._get_summarized_content_with_excerpts(documents, query)
                processing_note = f"Analyzed all {total_files} files using intelligent summarization"
            
            if not content_chunks:
                error_msg = "Unable to extract relevant content from the documents."
                return (error_msg, "") if include_context else error_msg
            
            # Build context for transparency
            context = f"Processing Strategy: {processing_note}\n" + "="*60 + "\n"
            context += "\n".join([f"File {i+1}: {chunk['filename']}" for i, chunk in enumerate(content_chunks[:10])])
            if len(content_chunks) > 10:
                context += f"\n... and {len(content_chunks) - 10} more files"
            
            # Generate flexible, user-centric response
            flexible_prompt = f"""
User Question: {query}

Document Content from {total_files} files:
{self._format_content_for_analysis(content_chunks)}

Please provide a detailed, conversational response to the user's question using information from all the files. Be thorough and comprehensive, including specific details, quotes, and context. Avoid structured formats with bullet points or section headers unless the user specifically requests that format.

Write as if you're having a natural conversation with the user, providing rich detail and specific information. Include relevant background context, specific quotes or examples, and elaborate on important points.

IMPORTANT: When referencing information, always cite the specific document filename rather than document numbers or chunk references. This helps users identify the source document.
"""

            messages = [
                SystemMessage(content=f"You are an expert analyst with access to {total_files} meeting documents. Provide detailed, conversational responses that avoid structured formats unless specifically requested. Include rich details, quotes, and context in your comprehensive answers. Always cite document filenames rather than document numbers when referencing information."),
                HumanMessage(content=flexible_prompt)
            ]
            
            response = self.llm.invoke(messages)
            
            # Add transparency note
            final_response = f"*Based on analysis of all {total_files} files in your project*\n\n{response.content}"
            
            return (final_response, context) if include_context else final_response
            
        except Exception as e:
            logger.error(f"Error in flexible comprehensive processing: {e}")
            error_msg = f"Error analyzing {total_files} files: {str(e)}"
            return (error_msg, "") if include_context else error_msg

    def _get_detailed_content_from_all_files(self, documents: List[Any]) -> List[Dict[str, str]]:
        """Get detailed content from all files for small projects (≤15 files)"""
        content_chunks = []
        for doc in documents:
            doc_info = self.vector_db.get_document_metadata(doc['document_id'])
            if doc_info:
                content_chunks.append({
                    'filename': doc_info['filename'],
                    'date': doc_info['date'][:10],
                    'content': doc_info['content_summary'],
                    'topics': ', '.join(doc_info['main_topics'][:3]),  # Top 3 topics
                    'participants': ', '.join(doc_info['participants'][:5])  # Top 5 participants
                })
        return content_chunks

    def _get_smart_sampled_content(self, documents: List[Any], query: str) -> List[Dict[str, str]]:
        """Get smart sampled content for medium projects (16-50 files)"""
        # For medium projects, get summaries from all files but prioritize based on relevance
        all_content = self._get_detailed_content_from_all_files(documents)
        
        # Simple relevance scoring based on query keywords
        query_words = set(query.lower().split())
        
        for chunk in all_content:
            # Score based on how many query words appear in content
            content_words = set(chunk['content'].lower().split())
            chunk['relevance_score'] = len(query_words.intersection(content_words))
        
        # Sort by relevance but keep all files
        all_content.sort(key=lambda x: x['relevance_score'], reverse=True)
        return all_content

    def _get_summarized_content_with_excerpts(self, documents: List[Any], query: str) -> List[Dict[str, str]]:
        """Get summarized content with key excerpts for large projects (50+ files)"""
        # For large projects, group by time periods and get summaries
        content_chunks = []
        
        # Group documents by month for better organization
        monthly_groups = {}
        for doc in documents:
            doc_info = self.vector_db.get_document_metadata(doc['document_id'])
            if doc_info:
                try:
                    import datetime
                    doc_date = datetime.datetime.fromisoformat(doc_info['date'].replace('Z', '+00:00'))
                    month_key = doc_date.strftime('%Y-%m')
                    if month_key not in monthly_groups:
                        monthly_groups[month_key] = []
                    monthly_groups[month_key].append(doc_info)
                except:
                    if 'unknown' not in monthly_groups:
                        monthly_groups['unknown'] = []
                    monthly_groups['unknown'].append(doc_info)
        
        # Create summaries for each time period
        for period, docs in monthly_groups.items():
            period_summary = f"Period {period} ({len(docs)} files): "
            topics = set()
            participants = set()
            
            for doc in docs:
                topics.update(doc['main_topics'][:2])  # Top 2 topics per doc
                participants.update(doc['participants'][:3])  # Top 3 participants per doc
            
            period_summary += f"Main topics: {', '.join(list(topics)[:5])}. "
            period_summary += f"Key participants: {', '.join(list(participants)[:8])}."
            
            content_chunks.append({
                'filename': f"{len(docs)} files from {period}",
                'date': period,
                'content': period_summary,
                'topics': ', '.join(list(topics)[:5]),
                'participants': ', '.join(list(participants)[:8])
            })
        
        return content_chunks

    def _format_content_for_analysis(self, content_chunks: List[Dict[str, str]]) -> str:
        """Format content chunks for AI analysis"""
        formatted_content = []
        
        for i, chunk in enumerate(content_chunks, 1):
            formatted_content.append(f"""
Document: {chunk['filename']} ({chunk['date']})

{chunk['content']}

{'='*60}""")
        
        return "\n".join(formatted_content)

    def process_file_async(self, file_path: str, filename: str, user_id: str,
                          project_id: str = None, meeting_id: str = None,
                          job_id: str = None, status_id: str = None) -> Dict:
        """Process a single file asynchronously"""
        try:
            # Update status
            if status_id:
                self.vector_db.update_file_processing_status(status_id, 'processing')
            
            # Calculate file hash
            file_hash = self.vector_db.calculate_file_hash(file_path)
            file_size = os.path.getsize(file_path)
            
            # Check for duplicates again (safety check)
            duplicate_info = self.vector_db.is_file_duplicate(file_hash, filename, user_id)
            if duplicate_info:
                if status_id:
                    self.vector_db.update_file_processing_status(status_id, 'skipped', 
                                                               f"Duplicate of {duplicate_info['original_filename']}")
                return {
                    'success': False,
                    'error': f'Duplicate file: {duplicate_info["original_filename"]}',
                    'is_duplicate': True,
                    'duplicate_info': duplicate_info
                }
            
            # Extract content
            content = self.read_document_content(file_path)
            if not content.strip():
                if status_id:
                    self.vector_db.update_file_processing_status(status_id, 'failed', 'Empty file content')
                return {'success': False, 'error': 'Empty file content'}
            
            # Store file metadata
            document_id = self.vector_db.store_document_metadata(
                filename, content, user_id, project_id, meeting_id
            )
            
            # Store file hash for deduplication
            hash_id = self.vector_db.store_file_hash(
                file_hash, filename, filename, file_size, user_id,
                project_id, meeting_id, document_id
            )
            
            # Extract actual meeting date from filename
            extracted_date = self.extract_date_from_filename(filename, content)
            
            # Create MeetingDocument object for processing
            meeting_doc = MeetingDocument(
                document_id=document_id,
                filename=filename,
                date=extracted_date,
                title=filename,
                content=content,
                content_summary="",
                main_topics=[],
                past_events=[],
                future_actions=[],
                participants=[],
                user_id=user_id,
                meeting_id=meeting_id,
                project_id=project_id,
                file_size=file_size
            )
            
            # Process with enhanced chunking
            chunks = self.chunk_document(meeting_doc)
            
            # Store document and chunks together using the existing add_document method
            self.vector_db.add_document(meeting_doc, chunks)
            
            # Update final status
            if status_id:
                self.vector_db.update_file_processing_status(status_id, 'completed', 
                                                           f'Successfully processed {len(chunks)} chunks')
            
            logger.info(f"Successfully processed {filename} with {len(chunks)} chunks")
            return {
                'success': True,
                'document_id': document_id,
                'chunks_created': len(chunks),
                'hash_id': hash_id
            }
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            if status_id:
                self.vector_db.update_file_processing_status(status_id, 'failed', str(e))
            return {'success': False, 'error': str(e)}

    def process_files_batch_async(self, files: List[Dict], user_id: str, 
                                 project_id: str = None, meeting_id: str = None, 
                                 max_workers: int = 3, job_id: str = None) -> Dict:
        """Process multiple files asynchronously with threading"""
        if not files:
            return {'success': False, 'error': 'No files provided'}
        
        # Create upload job if not provided
        if job_id is None:
            job_id = self.vector_db.create_upload_job(user_id, len(files), project_id, meeting_id)
        
        # Create file processing status entries
        file_tasks = []
        for file_info in files:
            file_hash = self.vector_db.calculate_file_hash(file_info['path'])
            file_size = os.path.getsize(file_info['path'])
            
            # Check for duplicates first
            duplicate_info = self.vector_db.is_file_duplicate(file_hash, file_info['filename'], user_id)
            if duplicate_info:
                logger.info(f"Skipping duplicate file: {file_info['filename']}")
                continue
            
            status_id = self.vector_db.create_file_processing_status(job_id, file_info['filename'], 
                                                         file_size, file_hash)
            file_tasks.append({
                'file_path': file_info['path'],
                'filename': file_info['filename'],
                'status_id': status_id,
                'file_hash': file_hash
            })
        
        if not file_tasks:
            self.vector_db.update_job_status(job_id, 'completed', 0, 0)
            return {
                'success': True,
                'job_id': job_id,
                'message': 'All files were duplicates, no processing needed'
            }
        
        # Update job with actual file count after deduplication
        self.vector_db.update_job_status(job_id, 'processing')
        
        # Process files concurrently
        results = []
        processed_count = 0
        failed_count = 0
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks
            future_to_task = {
                executor.submit(
                    self.process_file_async,
                    task['file_path'],
                    task['filename'],
                    user_id,
                    project_id,
                    meeting_id,
                    job_id,
                    task['status_id']
                ): task for task in file_tasks
            }
            
            # Process completed tasks
            for future in as_completed(future_to_task):
                task = future_to_task[future]
                try:
                    result = future.result()
                    results.append({
                        'filename': task['filename'],
                        'result': result
                    })
                    
                    if result['success']:
                        processed_count += 1
                    else:
                        failed_count += 1
                        
                    # Update job progress
                    self.vector_db.update_job_status(job_id, 'processing', processed_count, failed_count)
                    
                except Exception as e:
                    logger.error(f"Error processing {task['filename']}: {e}")
                    failed_count += 1
                    results.append({
                        'filename': task['filename'],
                        'result': {'success': False, 'error': str(e)}
                    })
                    self.vector_db.update_job_status(job_id, 'processing', processed_count, failed_count)
        
        # Update final job status
        final_status = 'completed' if failed_count == 0 else 'partial' if processed_count > 0 else 'failed'
        self.vector_db.update_job_status(job_id, final_status, processed_count, failed_count)
        
        # Save vector index if any files were processed successfully
        if processed_count > 0:
            self.vector_db.save_index()
            logger.info(f"Saved FAISS index after processing {processed_count} documents")
        
        logger.info(f"Batch processing completed. Job: {job_id}, "
                   f"Processed: {processed_count}, Failed: {failed_count}")
        
        return {
            'success': True,
            'job_id': job_id,
            'processed_files': processed_count,
            'failed_files': failed_count,
            'total_files': len(file_tasks),
            'results': results
        }

def main():
    """Main function for Meeting Document AI System with OpenAI"""
    
    pass
    
    try:
        # Check OpenAI API key
        if not openai_api_key:
            pass
            return
        
        pass
        
        processor = EnhancedMeetingDocumentProcessor(chunk_size=1000, chunk_overlap=200)
        
        # Check if documents folder exists
        docs_folder = Path("meeting_documents")
        if not docs_folder.exists():
            docs_folder.mkdir(exist_ok=True)
            print("📁 Created 'meeting_documents' folder. Please add your meeting documents to this folder.")
            return
        
        # Check for documents
        doc_files = []
        for ext in ['.txt', '.docx', '.pdf']:
            doc_files.extend(docs_folder.glob(f"*{ext}"))
        
        if not doc_files:
            print("📁 No meeting documents found in the 'meeting_documents' folder.")
            print("📋 Supported formats: .txt, .docx, .pdf")
            print("📂 Please add your meeting documents to the 'meeting_documents' folder and run again.")
            return
        
        print(f"📄 Found {len(doc_files)} documents to process")
        
        # Check if vector database exists
        print("🔍 Checking existing vector database...")
        if processor.vector_db.index.ntotal == 0:
            print("🔄 Processing documents and building vector database...")
            processor.process_documents("meeting_documents")
        else:
            print(f"[OK] Loaded existing vector database with {processor.vector_db.index.ntotal} vectors")
        
        # Show comprehensive statistics
        print("\n[STATS] System Statistics:")
        print("-" * 50)
        stats = processor.get_meeting_statistics()
        if "error" not in stats:
            print(f"📋 Total meetings: {stats.get('total_meetings', 0)}")
            print(f"🧩 Total chunks: {stats.get('total_chunks', 0)}")
            print(f"🔢 Vector index size: {stats.get('vector_index_size', 0)}")
            print(f"📏 Average chunk length: {stats.get('average_chunk_length', 0)} characters")
            
            if 'date_range' in stats:
                print(f"📅 Date range: {stats['date_range']['earliest']} to {stats['date_range']['latest']}")
            
            # AI Configuration details removed from display
        
        # Interactive query loop
        print("\n[RESULT] Interactive Query Session")
        print("-" * 50)
        print("💡 Now supports hundreds of documents with hybrid search!")
        print("🔍 Combines semantic similarity + keyword matching for better results")
        
        example_queries = [
            "What are the main topics from recent meetings?",
            "Tell me about the AI integration progress across all meetings",
            "What are our upcoming deadlines and action items?",
            "Summarize all migration plans discussed"
        ]
        
        for i, example in enumerate(example_queries[:3], 1):
            print(f"   {i}. {example}")
        print("   ... or ask anything about your meetings!")
        
        print(f"\n💬 Commands: 'quit'/'exit' to stop, 'stats' for statistics, 'help' for examples")
        
        while True:
            print("\n" + "-"*60)
            query = input("🤔 Your question: ").strip()
            
            if query.lower() in ['quit', 'exit', 'q']:
                print("👋 Thank you for using the Meeting Document AI System!")
                break
            elif query.lower() == 'stats':
                stats = processor.get_meeting_statistics()
                print("\n[STATS] Detailed Statistics:")
                print(json.dumps(stats, indent=2, default=str))
                continue
            elif query.lower() == 'help':
                print("\n📚 Example Queries (enhanced with hybrid search):")
                for i, example in enumerate(example_queries, 1):
                    print(f"   {i}. {example}")
                print("\n💭 More ideas:")
                print("   • Search across hundreds of documents instantly")
                print("   • Find specific keywords + semantic meaning")
                print("   • Get comprehensive answers from multiple meetings")
                print("   • Timeline analysis across large document sets")
                continue
            elif query.lower() == 'refresh':
                try:
                    processor.refresh_clients()
                    print("[OK] OpenAI clients refreshed successfully")
                except Exception as e:
                    print(f"[ERROR] Failed to refresh clients: {e}")
                continue
            elif not query:
                print("❓ Please enter a valid question.")
                continue
            
            print(f"\n🔍 Processing with hybrid search: '{query}'...")
            start_time = datetime.now()
            
            try:
                answer = processor.answer_query(query)
                processing_time = (datetime.now() - start_time).total_seconds()
                
                print(f"\n💬 Answer (processed in {processing_time:.2f}s):")
                print("-" * 60)
                print(answer)
                
            except Exception as query_error:
                print(f"[ERROR] Error processing query: {query_error}")
                print("🔄 You can try rephrasing your question or check your OpenAI API key.")
    
    except Exception as e:
        logger.error(f"Critical error in main execution: {e}")
        print(f"[ERROR] Critical Error: {e}")
        print("🔧 Please check your OpenAI API key and configuration.")
        print("💡 Make sure you have set OPENAI_API_KEY environment variable")

if __name__ == "__main__":
    main()